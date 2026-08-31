import io
import json
import re
import time
import traceback
import uuid
from datetime import date, datetime
from pathlib import Path
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

from memo_parser import parse_memo_text
from qa_engine import answer_question


# =========================================================
# 팀원 일정표 모듈 - HTML/JavaScript 원본 내장
# =========================================================
ONBOARDING_SCHEDULE_HTML = r'''<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>후임자 온보딩 일과표 자동 생성기</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@500;600;700&family=IBM+Plex+Sans+KR:wght@400;500;600;700&family=IBM+Plex+Mono:wght@400;500&display=swap" rel="stylesheet">
<script src="https://cdnjs.cloudflare.com/ajax/libs/mammoth/1.6.0/mammoth.browser.min.js"></script>
<style>
  :root{
    --ink:#161A1F; --paper:#F2F3EF; --surface:#FFFFFF;
    --line:#DFE1DA; --line-strong:#C7CAC1;
    --accent:#0F6657; --accent-dark:#0B4A40; --accent-soft:#DCEFEA;
    --amber:#9A6B14; --amber-soft:#F6E9D2;
    --warn:#B14A16; --warn-soft:#F5E2D5;
    --muted:#6B7069;
    --shadow: 0 1px 2px rgba(20,20,15,0.04), 0 8px 24px rgba(20,20,15,0.05);
  }
  *{box-sizing:border-box;}
  html,body{margin:0;padding:0;}
  body{background:var(--paper); color:var(--ink); font-family:'IBM Plex Sans KR', sans-serif; line-height:1.55; -webkit-font-smoothing:antialiased;}
  .mono{font-family:'IBM Plex Mono', monospace;}
  header{border-bottom:1px solid var(--line); background:var(--surface); padding:22px 32px;}
  header .mark{width:34px;height:34px;border-radius:8px;background:var(--accent);display:inline-flex;align-items:center;justify-content:center;color:white;font-weight:700;font-size:14px;font-family:'Space Grotesk',sans-serif;margin-right:10px;vertical-align:middle;}
  header h1{font-size:18px;margin:0;font-weight:700;display:inline;vertical-align:middle;}
  header p{margin:6px 0 0;font-size:12.5px;color:var(--muted);}
  .demo-banner{max-width:1360px;margin:16px auto 0;padding:0 24px;}
  .demo-banner .inner{background:var(--amber-soft);border:1px solid #E3CE9E;color:var(--amber);border-radius:10px;padding:10px 14px;font-size:12.3px;font-weight:600;line-height:1.5;}
  .layout{display:grid; grid-template-columns:320px 1fr; gap:20px; padding:24px; max-width:1360px; margin:0 auto; align-items:start;}
  @media (max-width:980px){ .layout{grid-template-columns:1fr;} }
  .panel{background:var(--surface); border:1px solid var(--line); border-radius:14px; box-shadow:var(--shadow);}
  .panel-head{padding:16px 18px 12px; border-bottom:1px solid var(--line);}
  .panel-head h2{font-size:13px;margin:0;font-weight:700;text-transform:uppercase;letter-spacing:0.05em;color:var(--muted);}
  .panel-body{padding:14px 18px 18px;}
  .source-block{margin-bottom:14px;}
  .source-block:last-child{margin-bottom:0;}
  .source-label{font-size:12px;font-weight:600;margin-bottom:5px;}
  .source-label .req{font-size:10.5px;color:var(--muted);font-weight:400;}
  .dropzone{border:1.5px dashed var(--line-strong); border-radius:10px; padding:12px 8px; text-align:center; cursor:pointer; background:#FBFBF9;}
  .dropzone.drag{border-color:var(--accent); background:var(--accent-soft);}
  .dropzone .dz-title{font-size:11.8px;font-weight:600;}
  .dropzone .dz-sub{font-size:10.4px;color:var(--muted);margin-top:2px;}
  input[type=file]{display:none;}
  .file-list{margin-top:6px;display:flex;flex-direction:column;gap:4px;}
  .file-item{display:flex;align-items:center;gap:6px;font-size:11px;padding:5px 8px;border-radius:6px;background:#FBFBF9;border:1px solid var(--line);}
  .file-item .fname{flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;font-weight:500;}
  .file-item .fstatus{font-size:9.5px;color:var(--muted);font-family:'IBM Plex Mono',monospace;}
  .file-item .fstatus.ok{color:var(--accent-dark);}
  .file-item .fstatus.err{color:var(--warn);}
  .file-item .fremove{cursor:pointer;color:var(--muted);font-size:12px;padding:1px 3px;}
  .btn{appearance:none;border:none;cursor:pointer;font-family:'IBM Plex Sans KR',sans-serif;font-weight:600;font-size:13px;border-radius:8px;padding:9px 14px;}
  .btn:disabled{opacity:0.5;cursor:default;}
  .btn-primary{background:var(--accent); color:white;}
  .btn-primary:hover:not(:disabled){background:var(--accent-dark);}
  .btn-ghost{background:transparent;color:var(--accent-dark);border:1px solid var(--line-strong);}
  .btn-block{width:100%;text-align:center;}
  .row-btns{display:flex;gap:8px;margin-top:10px;flex-wrap:wrap;}
  .week-tabs{display:flex;gap:8px;margin-bottom:16px;flex-wrap:wrap;}
  .week-tab{border:1px solid var(--line-strong);border-radius:10px;padding:8px 14px;cursor:pointer;background:var(--surface);font-size:12.5px;font-weight:600;}
  .week-tab.active{border-color:var(--accent);background:var(--accent-soft);}
  .grid-wrap{overflow-x:auto;}
  table.sched{border-collapse:collapse;width:100%;min-width:760px;}
  table.sched th{font-size:11px;text-transform:uppercase;letter-spacing:0.04em;color:var(--muted);font-weight:700;padding:8px 6px;border-bottom:1px solid var(--line);text-align:left;}
  table.sched td.timecol{font-family:'IBM Plex Mono',monospace;font-size:10.5px;color:var(--muted);white-space:nowrap;padding:10px 8px 10px 0;vertical-align:top;}
  table.sched td.cell{border:1px solid var(--line);padding:0;vertical-align:top;width:19%;}
  .cell-inner{padding:8px 9px;min-height:58px;font-size:12px;}
  .cell-inner.event{border-left:3px solid #C9A24A; background:var(--amber-soft);}
  .cell-inner.task{border-left:3px solid var(--accent);}
  .cell-inner.empty{border-left:3px solid var(--line-strong); background:#FBFBF9; color:var(--muted); font-size:11px;}
  .cell-task{font-weight:600;line-height:1.4; outline:none; border-radius:5px; padding:1px 3px; margin:-1px -3px; cursor:text;}
  .cell-task:hover{background:rgba(15,102,87,0.06);}
  .cell-task:focus{background:var(--surface);box-shadow:0 0 0 2px var(--accent-soft);}
  .cell-reason{font-size:10.3px;color:var(--muted);margin-top:4px;line-height:1.4;}
  .cell-badge{display:inline-block;font-size:9px;font-weight:700;padding:1px 6px;border-radius:999px;margin-top:5px;font-family:'IBM Plex Mono',monospace;background:var(--line); color:var(--muted);}
  .cell-badge.pri-긴급, .cell-badge.pri-상{background:var(--warn-soft);color:var(--warn);}
  .cell-badge.pri-중{background:var(--accent-soft);color:var(--accent-dark);}
  .cell-badge.edited{background:#EDE4F5;color:#6A4C93;margin-left:4px;}
  tr.lunch td.cell .cell-inner{background:#F1F1EC;color:var(--muted);font-size:11px;text-align:center;border-left:3px solid var(--line-strong);}
  .overflow-box{margin-top:14px;padding:12px 14px;border-radius:10px;border:1px solid var(--line);background:#FBFBF9;font-size:12px;}
  .overflow-box h4{margin:0 0 6px;font-size:11.5px;color:var(--muted);text-transform:uppercase;letter-spacing:0.04em;}
  .overflow-box ul{margin:0;padding-left:18px;}
  .mail-box{margin-top:10px;}
  .mail-box summary{cursor:pointer;font-size:11.5px;color:var(--muted);font-weight:600;}
  .mail-box .mail-item{font-size:11.5px;padding:6px 0;border-bottom:1px dashed var(--line);}
  .hint{font-size:11.8px;color:var(--muted);margin-top:4px;}
  .error{color:var(--warn);font-size:12px;margin-top:8px;}
  .empty-msg{font-size:13px;color:var(--muted);padding:24px 0;text-align:center;}
  .legend{display:flex;gap:14px;flex-wrap:wrap;margin-bottom:10px;}
  .legend span{font-size:10.6px;color:var(--muted);display:flex;align-items:center;gap:5px;}
  .legend .sw{width:9px;height:9px;border-radius:3px;display:inline-block;}
  .legend .sw.event{background:#C9A24A;}
  .legend .sw.task{background:var(--accent);}
  .legend .sw.empty{background:var(--line-strong);}
  .role-toggle{display:flex;align-items:center;gap:8px;margin-bottom:12px;flex-wrap:wrap;}
  .role-toggle .rt-label{font-size:11.5px;color:var(--muted);font-weight:600;}
  .role-chip{font-size:12px;font-weight:600;padding:6px 12px;border-radius:999px;border:1px solid var(--line-strong);background:var(--surface);cursor:pointer;}
  .role-chip.active{background:var(--ink);color:white;border-color:var(--ink);}
  .role-chip.active.mentor{background:#6A4C93;border-color:#6A4C93;}
  .cell-badge.edited.employee{background:var(--accent-soft);color:var(--accent-dark);}
  .cell-badge.edited.mentor{background:#EDE4F5;color:#6A4C93;}
  .week-tab .wt-sub{display:block;font-size:9.5px;font-weight:500;color:var(--muted);margin-top:2px;}
</style>
</head>
<body>

<header>
  <div><span class="mark">S</span><h1>후임자 온보딩 일과표 자동 생성기</h1></div>
  <p>회의록 · 캘린더 · 메일 문서를 업로드하면, 문서에 담긴 실제 일정과 미리 준비된 업무 자료를 합쳐서 <b>1주차는 최대한 빼곡하게</b>, <b>2주차는 실제 날짜에 있는 일정만</b> 반영한 일과표(월~금)를 만들어줍니다. 생성 후 칸을 클릭해 바로 수정할 수 있어요.</p>
</header>

<div class="layout">

  <!-- LEFT -->
  <div class="panel">
    <div class="panel-head"><h2>문서 업로드</h2></div>
    <div class="panel-body">

      <div class="source-block">
        <div class="source-label">회의록 <span class="req">(여러 개, docx)</span></div>
        <div class="dropzone" data-target="meeting">
          <div class="dz-title">파일 업로드</div>
          <div class="dz-sub">.docx (여러 개 선택 가능)</div>
        </div>
        <input type="file" class="fileInput" data-target="meeting" accept=".docx" multiple>
        <div class="file-list" id="fileList-meeting"></div>
      </div>

      <div class="source-block">
        <div class="source-label">캘린더 문서 <span class="req">(06_인수인계_캘린더.docx)</span></div>
        <div class="dropzone" data-target="calendar">
          <div class="dz-title">파일 업로드</div>
          <div class="dz-sub">.docx</div>
        </div>
        <input type="file" class="fileInput" data-target="calendar" accept=".docx">
        <div class="file-list" id="fileList-calendar"></div>
      </div>

      <div class="source-block">
        <div class="source-label">메일 <span class="req">(여러 개, 참고용 표시)</span></div>
        <div class="dropzone" data-target="email">
          <div class="dz-title">파일 업로드</div>
          <div class="dz-sub">.docx (여러 개 선택 가능)</div>
        </div>
        <input type="file" class="fileInput" data-target="email" accept=".docx" multiple>
        <div class="file-list" id="fileList-email"></div>
      </div>

      <div class="row-btns">
        <button class="btn btn-primary btn-block" id="genBtn">일정표 생성</button>
      </div>
      <p class="hint">업무일정·프로젝트 현황·자산·연락망 내용은 도구 안에 이미 반영되어 있어서 따로 업로드하지 않아도 됩니다.</p>
    </div>
  </div>

  <!-- RIGHT -->
  <div class="panel">
    <div class="panel-head"><h2>1주차 · 2주차 일과표</h2></div>
    <div class="panel-body">

      <div class="legend">
        <span><span class="sw event"></span>회의록/캘린더의 고정 시간 일정</span>
        <span><span class="sw task"></span>문서에서 온 업무/개요 항목</span>
        <span><span class="sw empty"></span>여유 시간</span>
      </div>

      <div class="role-toggle">
        <span class="rt-label">지금 누가 수정하나요?</span>
        <div class="role-chip active" id="role-employee" data-role="employee">신입</div>
        <div class="role-chip" id="role-mentor" data-role="mentor">선임자</div>
      </div>
      <p class="hint" style="margin-top:-6px;">신입 또는 선임자 버튼을 누른 뒤 칸을 클릭해 수정하면, 누가 고쳤는지 칸 아래 배지로 표시됩니다.</p>

      <div id="weekTabs" class="week-tabs"></div>

      <div class="row-btns" style="margin-top:0;margin-bottom:12px;">
        <button class="btn btn-ghost" id="resetBtn" style="display:none;">이 주 수정 초기화</button>
      </div>

      <div id="gridArea">
        <div class="empty-msg">왼쪽에서 문서를 업로드하고 "일정표 생성"을 눌러보세요.</div>
      </div>

      <div id="overflowArea"></div>
      <div id="mailArea"></div>
      <div id="laterArea"></div>

      <div id="errorBox" class="error" style="display:none;"></div>
    </div>
  </div>

</div>

<script>
const SLOTS = ["09:00-10:00","10:00-12:00","13:00-15:00","15:00-17:00","17:00-18:00"];
const SLOT_RANGES = [[540,600],[600,720],[780,900],[900,1020],[1020,1080]]; // minutes from midnight

let files = { meeting:[], calendar:[], email:[] };
let fileIdCounter = 0;

let originalWeeks = null; // parsed, immutable
let workingWeeks = null;  // editable copy
let weekOrder = [];
let currentWeekKey = null;
let currentRole = 'employee'; // 'employee' | 'mentor'

const el = id => document.getElementById(id);
const escapeHtml = s => (s||'').replace(/[&<>"']/g, c => ({'&':'&amp;','<':'&lt;','>':'&gt;','"':'&quot;',"'":'&#39;'}[c]));

function showError(msg){ el('errorBox').style.display='block'; el('errorBox').textContent = msg; }
function clearError(){ el('errorBox').style.display='none'; }

/* ---------------- 파일 업로드 UI ---------------- */

function renderFileList(target){
  const wrap = el('fileList-'+target);
  wrap.innerHTML = '';
  files[target].forEach(d=>{
    const row = document.createElement('div');
    row.className = 'file-item';
    const statusText = d.status === 'loading' ? '읽는 중…' : d.status === 'ok' ? '완료' : '실패';
    const statusClass = d.status === 'ok' ? 'ok' : d.status === 'err' ? 'err' : '';
    row.innerHTML = `<span class="fname">${escapeHtml(d.name)}</span><span class="fstatus ${statusClass}">${statusText}</span><span class="fremove" data-id="${d.id}">✕</span>`;
    row.querySelector('.fremove').onclick = () => {
      files[target] = files[target].filter(x => x.id !== d.id);
      renderFileList(target);
    };
    wrap.appendChild(row);
  });
}

async function handleFiles(target, fileListRaw, single){
  const arr = Array.from(fileListRaw);
  if(single){ files[target] = []; }
  for(const file of arr){
    const doc = { id: ++fileIdCounter, name: file.name, status: 'loading', raw:null };
    files[target].push(doc);
    renderFileList(target);
    try{
      const buf = await file.arrayBuffer();
      const res = await mammoth.extractRawText({ arrayBuffer: buf });
      doc.raw = res.value;
      doc.status = 'ok';
    }catch(err){
      doc.status = 'err';
      showError(`${file.name} 처리 실패: ${err.message}`);
    }
    renderFileList(target);
  }
}

document.querySelectorAll('.dropzone').forEach(dz=>{
  const target = dz.dataset.target;
  const input = document.querySelector(`.fileInput[data-target="${target}"]`);
  const single = !input.multiple;
  dz.addEventListener('click', ()=> input.click());
  input.addEventListener('change', e=>{ handleFiles(target, e.target.files, single); input.value=''; });
  ['dragenter','dragover'].forEach(evt=> dz.addEventListener(evt, e=>{ e.preventDefault(); dz.classList.add('drag'); }));
  ['dragleave','drop'].forEach(evt=> dz.addEventListener(evt, e=>{ e.preventDefault(); dz.classList.remove('drag'); }));
  dz.addEventListener('drop', e=>{ if(e.dataTransfer.files.length) handleFiles(target, e.dataTransfer.files, single); });
});

document.querySelectorAll('.role-chip').forEach(chip=>{
  chip.addEventListener('click', ()=>{
    currentRole = chip.dataset.role;
    document.querySelectorAll('.role-chip').forEach(c=>c.classList.remove('active','mentor'));
    chip.classList.add('active');
    if(currentRole === 'mentor') chip.classList.add('mentor');
  });
});

/* ---------------- 날짜/시간 유틸 ---------------- */

function normalizeDate(v){
  if(v instanceof Date && !isNaN(v)){
    const y=v.getFullYear(), m=String(v.getMonth()+1).padStart(2,'0'), d=String(v.getDate()).padStart(2,'0');
    return `${y}-${m}-${d}`;
  }
  if(typeof v === 'string'){
    let m = v.match(/(\d{4})-(\d{1,2})-(\d{1,2})/);
    if(m) return `${m[1]}-${String(m[2]).padStart(2,'0')}-${String(m[3]).padStart(2,'0')}`;
    m = v.match(/^(\d{1,2})\/(\d{1,2})$/);
    if(m) return `2026-${String(m[1]).padStart(2,'0')}-${String(m[2]).padStart(2,'0')}`;
  }
  return null;
}

function timeToMinutes(hhmm){
  const m = (hhmm||'').match(/(\d{1,2}):(\d{2})/);
  if(!m) return null;
  return (+m[1])*60 + (+m[2]);
}

function slotIndexForTime(hhmm){
  const min = timeToMinutes(hhmm);
  if(min===null) return 0;
  for(let i=0;i<SLOT_RANGES.length;i++){
    if(min >= SLOT_RANGES[i][0] && min < SLOT_RANGES[i][1]) return i;
  }
  if(min < SLOT_RANGES[0][0]) return 0;
  return SLOT_RANGES.length-1;
}

function weekKeyOf(dateStr){
  const d = new Date(dateStr+'T00:00:00');
  const day = d.getDay(); // 0=Sun
  const diffToMon = (day===0? -6 : 1-day);
  const mon = new Date(d); mon.setDate(d.getDate()+diffToMon);
  return mon.toISOString().slice(0,10);
}
function weekLabelOf(weekKey){
  const mon = new Date(weekKey+'T00:00:00');
  const fri = new Date(mon); fri.setDate(mon.getDate()+4);
  const f = d => `${d.getMonth()+1}/${d.getDate()}`;
  return `${f(mon)} ~ ${f(fri)}`;
}
function weekdaysOf(weekKey){
  const mon = new Date(weekKey+'T00:00:00');
  const out = [];
  const names = ['월요일','화요일','수요일','목요일','금요일'];
  for(let i=0;i<5;i++){
    const d = new Date(mon); d.setDate(mon.getDate()+i);
    out.push({ date: d.toISOString().slice(0,10), name: names[i], label: `${d.getMonth()+1}/${d.getDate()}` });
  }
  return out;
}

/* ---------------- 내장 데이터 (01~04 엑셀 내용을 코드에 미리 반영) ---------------- */
// 회의록/캘린더/메일 문서만 업로드해도 동일한 결과가 나오도록,
// 01~04 엑셀에 있던 내용을 그대로 코드 안에 옮겨두었습니다.
// (실제 서비스라면 이 부분도 업로드/DB 연동으로 대체할 수 있습니다)

const EMBEDDED_SCHEDULE_TASKS = [
  { date:"2026-09-01", title:"최종 점검보고서 제출", priority:"긴급", nextAction:"누락 사진 3장 확인 → 팀장 검토 → 고객사 제출", note:"오전 중 제출 권장", source:"업무일정 엑셀" },
  { date:"2026-09-02", title:"감지기 교체 견적 회신", priority:"상", nextAction:"한빛전기 단가 회신 확인 후 견적서 최종 작성", note:"협력사 회신 대기", source:"업무일정 엑셀" },
  { date:"2026-09-03", title:"고객사 정기회의 참석", priority:"상", nextAction:"회의 전 미조치 2건 및 사진자료 정리", note:"회의 14:00", source:"업무일정 엑셀" },
  { date:"2026-09-04", title:"주간 미완료 업무 점검", priority:"중", nextAction:"미완료 보고서/견적/고객 요청 목록 업데이트", note:"매주 금요일 반복", source:"업무일정 엑셀" },
  { date:"2026-09-07", title:"공용드라이브 권한 이관 요청", priority:"상", nextAction:"팀장 승인 후 이서연 계정에 편집권한 부여 요청", note:"외부공유 권한 제외", source:"업무일정 엑셀" },
  { date:"2026-09-08", title:"월간 점검 일정 확정", priority:"중", nextAction:"현장팀 일정 취합 후 9월 점검표 확정", note:"현장팀 3명 일정 확인 필요", source:"업무일정 엑셀" },
  { date:"2026-09-10", title:"D물류센터 사전자료 요청", priority:"중", nextAction:"도면/설비목록/이전 점검결과 요청 메일 발송", note:"신규 인계 후 첫 신규 현장", source:"업무일정 엑셀" },
];

const EMBEDDED_PROJECT_DEADLINE_TASKS = [
  { date:"2026-09-01", title:"[마감] A동 소방시설 정기점검", priority:"상", nextAction:"누락 사진 3장 확인 후 최종 제출", note:"사진 누락 시 제출 지연 가능", source:"프로젝트 진행현황 엑셀" },
  { date:"2026-09-02", title:"[마감] B공장 감지기 교체", priority:"상", nextAction:"협력사 단가 반영 후 견적 회신", note:"단가 지연 시 고객 회신 지연", source:"프로젝트 진행현황 엑셀" },
  { date:"2026-09-03", title:"[마감] C센터 종합정밀점검", priority:"상", nextAction:"9/3 회의에서 미조치 2건 일정 확정", note:"미조치 2건 일정 미확정", source:"프로젝트 진행현황 엑셀" },
  { date:"2026-09-10", title:"[마감] D물류센터 신규점검", priority:"상", nextAction:"사전자료 요청 및 현장 일정 협의", note:"자료 미수신 시 현장 준비 지연", source:"프로젝트 진행현황 엑셀" },
];

const EMBEDDED_ASSET_DEADLINE_TASKS = [
  { date:"2026-09-07", title:"[마감] 공용드라이브 Z: 인계", priority:"상", nextAction:"팀장 승인 후 IT 요청", note:"외부 공유 권한은 부여하지 않음", source:"계정·권한·자산 엑셀" },
  { date:"2026-09-04", title:"[마감] 고객요청 관리 엑셀 인계", priority:"상", nextAction:"최신 파일 경로 전달", note:"중복본 사용 금지", source:"계정·권한·자산 엑셀" },
  { date:"2026-09-07", title:"[마감] A동 현장 폴더 인계", priority:"상", nextAction:"공용드라이브 권한과 함께 이관", note:"보고서 최종본 폴더 확인", source:"계정·권한·자산 엑셀" },
  { date:"2026-09-04", title:"[마감] 법인 태블릿 2번 인계", priority:"상", nextAction:"자산대장 서명 후 인계", note:"충전기 포함", source:"계정·권한·자산 엑셀" },
];

const EMBEDDED_PROJECT_OVERVIEW = [
  { title:"[개요] A동 소방시설 정기점검 현황 파악", detail:"보고서 최종화 · 진행률 80% · 다음액션: 누락 사진 3장 확인 후 최종 제출", badge:"개요", source:"프로젝트 진행현황 엑셀" },
  { title:"[개요] B공장 감지기 교체 현황 파악", detail:"견적 작성 · 진행률 60% · 다음액션: 협력사 단가 반영 후 견적 회신", badge:"개요", source:"프로젝트 진행현황 엑셀" },
  { title:"[개요] C센터 종합정밀점검 현황 파악", detail:"후속조치 협의 · 진행률 30% · 다음액션: 9/3 회의에서 미조치 2건 일정 확정", badge:"개요", source:"프로젝트 진행현황 엑셀" },
  { title:"[개요] D물류센터 신규점검 현황 파악", detail:"사전준비 · 진행률 10% · 다음액션: 사전자료 요청 및 현장 일정 협의", badge:"개요", source:"프로젝트 진행현황 엑셀" },
];

const EMBEDDED_ASSET_OVERVIEW = [
  { title:"[개요] 공용드라이브 Z: 인계 상태 확인", detail:"상태: 인계예정 · 인계방법: 팀장 승인 후 IT 요청 · 주의사항: 외부 공유 권한은 부여하지 않음", badge:"개요", source:"계정·권한·자산 엑셀" },
  { title:"[개요] 고객요청 관리 엑셀 인계 상태 확인", detail:"상태: 미완료 · 인계방법: 최신 파일 경로 전달 · 주의사항: 중복본 사용 금지", badge:"개요", source:"계정·권한·자산 엑셀" },
  { title:"[개요] 사내 IT요청 포털 인계 상태 확인", detail:"상태: 완료 · 인계방법: 개인 계정 직접 로그인 · 주의사항: 비밀번호 공유 금지", badge:"개요", source:"계정·권한·자산 엑셀" },
  { title:"[개요] A동 현장 폴더 인계 상태 확인", detail:"상태: 인계예정 · 인계방법: 공용드라이브 권한과 함께 이관 · 주의사항: 보고서 최종본 폴더 확인", badge:"개요", source:"계정·권한·자산 엑셀" },
  { title:"[개요] 법인 태블릿 2번 인계 상태 확인", detail:"상태: 미완료 · 인계방법: 자산대장 서명 후 인계 · 주의사항: 충전기 포함", badge:"개요", source:"계정·권한·자산 엑셀" },
];

const EMBEDDED_CONTACTS = [
  { title:"박현우(부장, 세림관리) 컨택포인트 파악", detail:"관련 업무: A동 점검보고서 · 유의사항: 보고서 전달 전 전화로 먼저 안내. 오전 11시 이전 연락 선호.", badge:"소개", source:"담당자 연락망 엑셀" },
  { title:"최은지(대리, B공장 시설팀) 컨택포인트 파악", detail:"관련 업무: B공장 견적/교체 일정 · 유의사항: 메일 제목에 [B공장] 표기. 견적 수정사항은 표로 정리.", badge:"소개", source:"담당자 연락망 엑셀" },
  { title:"조민석(과장, 한빛전기) 컨택포인트 파악", detail:"관련 업무: 감지기 단가/납기 · 유의사항: 급한 건 전화, 일반 단가 문의는 문자 가능.", badge:"소개", source:"담당자 연락망 엑셀" },
  { title:"오세훈(과장, C센터 시설팀) 컨택포인트 파악", detail:"관련 업무: C센터 후속조치 · 유의사항: 회의자료는 전날 17시까지 공유.", badge:"소개", source:"담당자 연락망 엑셀" },
  { title:"박지훈(대리, 현장점검팀) 컨택포인트 파악", detail:"관련 업무: 현장 사진/점검결과 · 유의사항: 사진 누락 확인은 박지훈 대리에게 요청.", badge:"소개", source:"담당자 연락망 엑셀" },
  { title:"한유리(사원, 영업팀) 컨택포인트 파악", detail:"관련 업무: 견적/고객 요청 · 유의사항: 견적 금액 변경 시 반드시 공유.", badge:"소개", source:"담당자 연락망 엑셀" },
];

/* ---------------- DOCX(텍스트) 파싱 ---------------- */

function splitLines(rawText){
  return rawText.split('\n').map(s=>s.trim()).filter(Boolean);
}

// 캘린더 문서: 표 헤더 "날짜","시간","일정","내용" 다음에 4개씩 데이터가 이어지는 구조
function parseCalendarDocx(rawText){
  const lines = splitLines(rawText);
  let start = -1;
  for(let i=0;i<lines.length-3;i++){
    if(lines[i]==='날짜' && lines[i+1]==='시간' && lines[i+2]==='일정' && lines[i+3]==='내용'){ start = i+4; break; }
  }
  const events = [];
  if(start<0) return events;
  for(let i=start;i+3<lines.length;i+=4){
    const [dateStr, timeStr, title, detail] = lines.slice(i,i+4);
    const date = normalizeDate(dateStr);
    if(!date) break;
    const tm = timeStr.match(/(\d{1,2}:\d{2})\s*-\s*(\d{1,2}:\d{2})/);
    events.push({
      date,
      start: tm ? tm[1] : null,
      end: tm ? tm[2] : null,
      allDay: !tm,
      title,
      detail,
      source: '캘린더 문서',
    });
  }
  return events;
}

// 회의록 문서: "일시" 다음줄에 "YYYY-MM-DD HH:MM-HH:MM", 첫 줄이 제목
function parseMeetingDocx(rawText, filename){
  const lines = splitLines(rawText);
  if(!lines.length) return null;
  const title = lines[0];
  let dtLine = null, place='', attendees='';
  for(let i=0;i<lines.length;i++){
    if(lines[i]==='일시' && lines[i+1]) dtLine = lines[i+1];
    if(lines[i]==='장소' && lines[i+1]) place = lines[i+1];
    if(lines[i]==='참석자' && lines[i+1]) attendees = lines[i+1];
  }
  if(!dtLine) return null;
  const m = dtLine.match(/(\d{4}-\d{2}-\d{2})\s+(\d{1,2}:\d{2})\s*-\s*(\d{1,2}:\d{2})/);
  if(!m) return null;
  return {
    date: normalizeDate(m[1]),
    start: m[2], end: m[3],
    allDay:false,
    title,
    detail: `참석: ${attendees}${place? ' · 장소: '+place : ''}`,
    source: '회의록 (' + filename + ')',
  };
}

// 메일 문서: 참고용 — 제목/받는사람/일시만 추출
function parseEmailDocx(rawText, filename){
  const lines = splitLines(rawText);
  let subject='', to='', dt='';
  for(let i=0;i<lines.length;i++){
    if(lines[i]==='제목' && lines[i+1]) subject = lines[i+1];
    if(lines[i]==='받는사람' && lines[i+1]) to = lines[i+1];
    if(lines[i]==='일시' && lines[i+1]) dt = lines[i+1];
  }
  const date = normalizeDate(dt);
  if(!subject) return null;
  return { date, subject, to, raw: dt, source: filename };
}

/* ---------------- 일정표 생성 ---------------- */

function buildSchedule(){
  const scheduleTasks = [ ...EMBEDDED_SCHEDULE_TASKS, ...EMBEDDED_PROJECT_DEADLINE_TASKS, ...EMBEDDED_ASSET_DEADLINE_TASKS ];
  const timedEvents = [];
  const mailRefs = [];

  files.calendar.filter(f=>f.status==='ok').forEach(f => {
    parseCalendarDocx(f.raw).forEach(ev => {
      if(ev.allDay){
        scheduleTasks.push({ date:ev.date, title:ev.title, project:'', owner:'', status:'', priority:'상', nextAction: ev.detail, note:'', source: ev.source });
      } else {
        timedEvents.push(ev);
      }
    });
  });

  files.meeting.filter(f=>f.status==='ok').forEach(f => {
    const ev = parseMeetingDocx(f.raw, f.name);
    if(ev) timedEvents.push(ev);
  });

  files.email.filter(f=>f.status==='ok').forEach(f => {
    const m = parseEmailDocx(f.raw, f.name);
    if(m) mailRefs.push(m);
  });

  if(scheduleTasks.length===0 && timedEvents.length===0){
    return null;
  }

  const priorityRank = { '긴급':0, '상':1, '중':2, '하':3, '':2 };
  scheduleTasks.sort((a,b) => (priorityRank[a.priority]??2) - (priorityRank[b.priority]??2));

  // 같은 날짜+같은 시작시간의 일정은 캘린더/회의록 등 여러 문서에 중복 기록된 것으로 보고 하나만 사용
  const seenEventKeys = new Set();
  const dedupedEvents = [];
  timedEvents.forEach(ev => {
    const key = ev.date + '|' + (ev.start || 'allday');
    if(seenEventKeys.has(key)) return;
    seenEventKeys.add(key);
    dedupedEvents.push(ev);
  });

  const tasksByDate = {};
  scheduleTasks.forEach(t => { (tasksByDate[t.date] = tasksByDate[t.date]||[]).push(t); });
  const eventsByDate = {};
  dedupedEvents.forEach(e => { (eventsByDate[e.date] = eventsByDate[e.date]||[]).push(e); });
  const mailByDate = {};
  mailRefs.forEach(m => { if(m.date) (mailByDate[m.date] = mailByDate[m.date]||[]).push(m); });

  const allDates = Array.from(new Set([...Object.keys(tasksByDate), ...Object.keys(eventsByDate)])).sort();

  // 실제 달력 날짜가 아니라 "1주차/2주차" 온보딩 개념이므로, 문서에서 발견된 날짜를
  // 시간순으로 최대 2개 주(캘린더 주 단위)까지만 사용하고 나머지는 별도로 남겨둔다.
  const weekKeysInOrder = Array.from(new Set(allDates.map(weekKeyOf))).sort();
  const usedWeekKeys = weekKeysInOrder.slice(0, 2);
  const week1Key = usedWeekKeys[0];
  const week2Key = usedWeekKeys[1];
  const laterDates = allDates.filter(d => !usedWeekKeys.includes(weekKeyOf(d)));

  const DAY_NAMES = ['월','화','수','목','금'];
  const weeks = {};

  function emptySlots(){
    return SLOTS.map(s => ({ slot:s, type:'empty', task:'', detail:'', badge:'', source:'', edited:false, editedBy:null }));
  }
  function placeIntoDay(dayObj, item){
    const freeIdx = dayObj.slots.findIndex(s => s.type==='empty');
    if(freeIdx>=0){
      dayObj.slots[freeIdx] = { slot:SLOTS[freeIdx], type:'task', task:item.title, detail:item.detail||'', badge:item.badge||'', source:item.source||'', edited:false, editedBy:null };
      return true;
    }
    dayObj.overflow.push(`${item.title} — ${item.source||''}`);
    return false;
  }

  // ---------- 2주차 (있는 경우): 실제 날짜에 있는 항목만 반영, 빈 칸은 그대로 둠 ----------
  if(week2Key){
    weeks[week2Key] = { label: '2주차', days:{}, mail:{} };
    allDates.forEach(date => {
      if(weekKeyOf(date) !== week2Key) return;
      const d = new Date(date+'T00:00:00');
      const dayIdx = d.getDay() - 1;
      if(dayIdx < 0 || dayIdx > 4) return;
      const dayName = DAY_NAMES[dayIdx];
      if(!weeks[week2Key].days[dayName]) weeks[week2Key].days[dayName] = { slots: emptySlots(), overflow: [] };
      const dayObj = weeks[week2Key].days[dayName];

      (eventsByDate[date]||[]).forEach(ev => {
        const idx = slotIndexForTime(ev.start);
        if(dayObj.slots[idx].type==='empty'){
          dayObj.slots[idx] = { slot:SLOTS[idx], type:'event', task:`${ev.start}-${ev.end} ${ev.title}`, detail: ev.detail||'', badge:'', source: ev.source, edited:false, editedBy:null };
        } else {
          dayObj.overflow.push(`(추가 일정) ${ev.start}-${ev.end} ${ev.title} — ${ev.source}`);
        }
      });
      (tasksByDate[date]||[]).forEach(t => {
        placeIntoDay(dayObj, { title:t.title, detail: t.nextAction ? `다음 조치: ${t.nextAction}` : (t.note||''), badge:t.priority, source:t.source });
      });
      if(mailByDate[date]) weeks[week2Key].mail[dayName] = (weeks[week2Key].mail[dayName]||[]).concat(mailByDate[date]);
    });
  }

  // ---------- 1주차: 업로드된 모든 문서 내용을 최대한 빼곡하게 채운 "초안" ----------
  if(week1Key){
    weeks[week1Key] = { label: '1주차', days:{}, mail:{} };
    DAY_NAMES.forEach(dn => { weeks[week1Key].days[dn] = { slots: emptySlots(), overflow: [] }; });

    // 1) 실제로 시간이 정해진 회의/캘린더 일정은 그대로 고정 배치 (진짜 약속이므로)
    allDates.forEach(date => {
      if(weekKeyOf(date) !== week1Key) return;
      const d = new Date(date+'T00:00:00');
      const dayIdx = d.getDay() - 1;
      if(dayIdx < 0 || dayIdx > 4) return;
      const dayName = DAY_NAMES[dayIdx];
      const dayObj = weeks[week1Key].days[dayName];
      (eventsByDate[date]||[]).forEach(ev => {
        const idx = slotIndexForTime(ev.start);
        if(dayObj.slots[idx].type==='empty'){
          dayObj.slots[idx] = { slot:SLOTS[idx], type:'event', task:`${ev.start}-${ev.end} ${ev.title}`, detail: ev.detail||'', badge:'', source: ev.source, edited:false, editedBy:null };
        }
      });
      if(mailByDate[date]) weeks[week1Key].mail[dayName] = (weeks[week1Key].mail[dayName]||[]).concat(mailByDate[date]);
    });

    // 2) 문서 전체 내용을 "학습 풀"로 모아서 남은 칸을 순서대로 채운다
    const pool = [];
    pool.push(...EMBEDDED_PROJECT_OVERVIEW);
    EMBEDDED_SCHEDULE_TASKS.forEach(t => pool.push({
      title: `실습: ${t.title}`,
      detail: t.nextAction ? `다음 조치: ${t.nextAction}` : (t.note||''),
      badge: '실습', source: t.source,
    }));
    pool.push(...EMBEDDED_ASSET_OVERVIEW);
    pool.push(...EMBEDDED_CONTACTS);
    mailRefs.forEach(m => pool.push({
      title: `메일 확인: ${m.subject}`,
      detail: `수신: ${m.to}${m.raw ? ' · '+m.raw : ''}`,
      badge: '메일', source: `메일 문서(${m.source})`,
    }));

    let pi = 0;
    DAY_NAMES.forEach(dn => {
      const dayObj = weeks[week1Key].days[dn];
      for(let si=0; si<SLOTS.length; si++){
        if(dayObj.slots[si].type !== 'empty') continue;
        if(pi >= pool.length) break;
        placeIntoDay(dayObj, pool[pi]);
        pi++;
      }
    });
    // 칸을 다 채우고도 남은 항목은 오늘 못 다룬 참고 자료로 남겨둔다
    if(pi < pool.length){
      weeks[week1Key].days['금'].overflow.push(...pool.slice(pi).map(p => `${p.title} — ${p.source||''}`));
    }
  }

  return { weeks, laterDates, weekOrder: [week1Key, week2Key].filter(Boolean) };
}

function deepCopy(obj){ return JSON.parse(JSON.stringify(obj)); }

function generate(){
  clearError();
  const result = buildSchedule();
  if(!result){
    showError('일정표를 만들지 못했습니다. 페이지를 새로고침한 뒤 다시 시도해주세요.');
    return;
  }
  originalWeeks = result.weeks;
  workingWeeks = deepCopy(result.weeks);
  weekOrder = result.weekOrder;
  currentWeekKey = weekOrder[0];
  renderWeekTabs();
  renderWeek(currentWeekKey);
  renderLaterNote(result.laterDates);
}

function renderLaterNote(laterDates){
  const host = el('laterArea');
  if(!laterDates || !laterDates.length){ host.innerHTML=''; return; }
  host.innerHTML = `<div class="overflow-box"><h4>1~2주차 이후로 밀린 일정 (${laterDates.length}건, 표에는 반영되지 않음)</h4><div class="hint">문서에 포함된 날짜 중 처음 2개 주차 범위를 벗어난 항목입니다. 3주차 일과표가 필요하면 알려주세요.</div></div>`;
}

function renderWeekTabs(){
  const wrap = el('weekTabs');
  wrap.innerHTML = '';
  weekOrder.forEach(wk => {
    const tab = document.createElement('div');
    tab.className = 'week-tab' + (wk===currentWeekKey ? ' active':'');
    tab.textContent = workingWeeks[wk].label;
    tab.onclick = () => { currentWeekKey = wk; renderWeekTabs(); renderWeek(wk); };
    wrap.appendChild(tab);
  });
}

const DAY_NAMES_FULL = { '월':'월요일', '화':'화요일', '수':'수요일', '목':'목요일', '금':'금요일' };

function renderWeek(wk){
  const weekData = workingWeeks[wk];
  const dayKeys = ['월','화','수','목','금'];
  el('resetBtn').style.display = 'inline-block';

  let html = '<div class="grid-wrap"><table class="sched"><thead><tr><th></th>';
  dayKeys.forEach(dk => html += `<th>${DAY_NAMES_FULL[dk]}</th>`);
  html += '</tr></thead><tbody>';

  SLOTS.forEach((slot, si) => {
    html += `<tr><td class="timecol">${slot}</td>`;
    dayKeys.forEach(dk => {
      const dayEntry = weekData.days[dk];
      const cellData = dayEntry ? dayEntry.slots[si] : { type:'empty', task:'', detail:'', badge:'', source:'', edited:false };
      const cls = cellData.type;
      html += `<td class="cell"><div class="cell-inner ${cls}" data-day="${dk}" data-slotidx="${si}">`;
      if(cellData.type==='empty'){
        html += `<div>여유 시간</div>`;
      } else {
        html += `<div class="cell-task" contenteditable="true" spellcheck="false">${escapeHtml(cellData.task)}</div>`;
        if(cellData.detail) html += `<div class="cell-reason">${escapeHtml(cellData.detail)}</div>`;
        if(cellData.badge) html += `<span class="cell-badge pri-${escapeHtml(cellData.badge)}">${escapeHtml(cellData.badge)}</span>`;
        if(cellData.editedBy) html += `<span class="cell-badge edited ${cellData.editedBy}">✎ ${cellData.editedBy==='mentor'?'선임자':'신입'} 수정</span>`;
        if(cellData.source) html += `<div class="cell-reason" style="opacity:0.7;">출처: ${escapeHtml(cellData.source)}</div>`;
      }
      html += `</div></td>`;
    });
    html += '</tr>';
    if(slot === '10:00-12:00'){
      html += `<tr class="lunch"><td class="timecol">12:00-13:00</td>`;
      dayKeys.forEach(()=> html += `<td class="cell"><div class="cell-inner">점심시간</div></td>`);
      html += `</tr>`;
    }
  });

  html += '</tbody></table></div>';
  el('gridArea').innerHTML = html;

  // overflow
  const overflowLines = [];
  dayKeys.forEach(dk => {
    const dayEntry = weekData.days[dk];
    if(dayEntry && dayEntry.overflow.length){
      dayEntry.overflow.forEach(line => overflowLines.push(`${DAY_NAMES_FULL[dk]}: ${line}`));
    }
  });
  el('overflowArea').innerHTML = overflowLines.length
    ? `<div class="overflow-box"><h4>이번 주 시간표에 다 못 들어간 항목</h4><ul>${overflowLines.map(l=>`<li>${escapeHtml(l)}</li>`).join('')}</ul></div>`
    : '';

  // mail refs
  const mailLines = [];
  dayKeys.forEach(dk => {
    (weekData.mail[dk]||[]).forEach(m => mailLines.push(`${DAY_NAMES_FULL[dk]} · ${m.raw||''} · ${m.subject} → ${m.to}`));
  });
  el('mailArea').innerHTML = mailLines.length
    ? `<details class="mail-box"><summary>참고: 이번 주 관련 메일 ${mailLines.length}건</summary>${mailLines.map(l=>`<div class="mail-item">${escapeHtml(l)}</div>`).join('')}</details>`
    : '';
}

el('gridArea').addEventListener('input', e=>{
  const target = e.target;
  if(!target.classList.contains('cell-task')) return;
  const cellEl = target.closest('.cell-inner');
  const dk = cellEl.dataset.day, slotIdx = +cellEl.dataset.slotidx;
  const dayEntry = workingWeeks[currentWeekKey].days[dk];
  if(!dayEntry) return;
  dayEntry.slots[slotIdx].task = target.textContent.trim();
  dayEntry.slots[slotIdx].edited = true;
  dayEntry.slots[slotIdx].editedBy = currentRole;
  let badge = cellEl.querySelector('.cell-badge.edited');
  if(!badge){
    badge = document.createElement('span');
    badge.className = 'cell-badge edited';
    cellEl.appendChild(badge);
  }
  badge.className = 'cell-badge edited ' + currentRole;
  badge.textContent = `✎ ${currentRole==='mentor'?'선임자':'신입'} 수정`;
});

el('resetBtn').addEventListener('click', ()=>{
  if(!originalWeeks || !currentWeekKey) return;
  workingWeeks[currentWeekKey] = deepCopy(originalWeeks[currentWeekKey]);
  renderWeek(currentWeekKey);
});

el('genBtn').addEventListener('click', generate);
</script>
</body>
</html>
'''



# =========================================================
# 0. Streamlit 기본 설정
# =========================================================
st.set_page_config(
    page_title="업무 인수인계 자동화",
    page_icon="📄",
    layout="wide",
)
st.title("📄 업무 인수인계 자동화 시스템")
st.caption(
    "API 없이 업무메모와 Excel 업무자료를 정리하고, "
    "표준 인수인계서(DOCX)를 생성하는 MVP"
)


# =========================================================
# 1. 자동 오류 감지 / 오류 리포트 / 사용자 행동 로그
# =========================================================
ERROR_REPORT_PATH = Path("data/error_reports.jsonl")
SLOW_RESPONSE_SECONDS = 8.0


def _mask_sensitive_text(value):
    """오류 신고 전에 주민등록번호처럼 명확한 민감 패턴을 간단히 마스킹."""
    if value is None:
        return ""

    text = str(value)
    text = re.sub(
        r"\b(\d{6})[- ]?([1-4])\d{6}\b",
        r"\1-*******",
        text,
    )
    return text


def _load_error_reports():
    if not ERROR_REPORT_PATH.exists():
        return []

    reports = []
    try:
        with ERROR_REPORT_PATH.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    reports.append(json.loads(line))
                except json.JSONDecodeError:
                    continue
    except OSError:
        return []

    return reports[-100:]


def init_error_reporting():
    defaults = {
        "error_action_logs": [],
        "last_screen": "🏠 대시보드",
        "last_function": "앱 이용",
        "last_input": "",
        "last_response_time": None,
        "last_error": None,
        "last_uploaded_files": [],
        "pending_auto_error_dialog": False,
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    if "error_reports" not in st.session_state:
        st.session_state["error_reports"] = _load_error_reports()


def track_action(
    screen,
    function,
    action,
    input_value="",
    metadata=None,
):
    """최근 사용자 행동을 최대 20개까지 세션에 기록."""
    event = {
        "time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "screen": screen,
        "function": function,
        "action": action,
        "input": _mask_sensitive_text(input_value),
        "metadata": metadata or {},
    }

    logs = list(st.session_state.get("error_action_logs", []))
    logs.append(event)
    st.session_state["error_action_logs"] = logs[-20:]
    st.session_state["last_screen"] = screen
    st.session_state["last_function"] = function

    if input_value:
        st.session_state["last_input"] = _mask_sensitive_text(input_value)


def _queue_auto_error_dialog():
    """오류 발생 후 사용자가 별도 버튼을 누르지 않아도 신고 팝업을 열도록 예약."""
    st.session_state["pending_auto_error_dialog"] = True


def record_error(
    screen,
    function,
    error,
    code="APP_ERROR",
    input_value="",
    response_time=None,
):
    """Python 예외를 자동 감지하고 오류 리포트 초안을 준비."""
    error_info = {
        "code": code,
        "type": type(error).__name__,
        "message": _mask_sensitive_text(str(error)),
        "occurred_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "response_time": response_time,
        "traceback": traceback.format_exc(),
    }

    st.session_state["last_error"] = error_info
    st.session_state["last_response_time"] = response_time

    track_action(
        screen=screen,
        function=function,
        action=f"오류 자동 감지 ({code})",
        input_value=input_value,
        metadata={
            "error_type": error_info["type"],
            "error_message": error_info["message"],
        },
    )
    _queue_auto_error_dialog()


def record_system_issue(
    screen,
    function,
    code,
    message,
    issue_type="SYSTEM_WARNING",
    input_value="",
    response_time=None,
):
    """예외가 없어도 응답 지연처럼 시스템이 판단 가능한 이상 상태를 자동 감지."""
    error_info = {
        "code": code,
        "type": issue_type,
        "message": _mask_sensitive_text(message),
        "occurred_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "response_time": response_time,
        "traceback": "",
    }

    st.session_state["last_error"] = error_info
    st.session_state["last_response_time"] = response_time

    track_action(
        screen=screen,
        function=function,
        action=f"이상 상태 자동 감지 ({code})",
        input_value=input_value,
        metadata={
            "error_type": issue_type,
            "error_message": error_info["message"],
        },
    )
    _queue_auto_error_dialog()


def build_error_report_draft():
    logs = st.session_state.get("error_action_logs", [])[-8:]
    last_error = st.session_state.get("last_error") or {}
    response_time = last_error.get(
        "response_time",
        st.session_state.get("last_response_time"),
    )

    action_lines = []
    for idx, event in enumerate(logs, start=1):
        line = (
            f"{idx}. [{event['time']}] {event['screen']} · "
            f"{event['action']}"
        )
        if event.get("input"):
            line += f" · 입력: {event['input']}"
        action_lines.append(line)

    error_code = last_error.get("code", "AUTO_DETECT")
    error_message = last_error.get("message", "자동 감지된 오류 메시지 없음")

    summary_parts = [
        f"{st.session_state.get('last_screen', '화면 미확인')}에서 ",
        f"{st.session_state.get('last_function', '기능 미확인')} 기능 이용 중 시스템이 오류를 자동 감지했습니다.",
    ]

    if st.session_state.get("last_input"):
        summary_parts.append(
            f" 최근 입력/검색어는 '{st.session_state['last_input']}'입니다."
        )

    if last_error:
        summary_parts.append(
            f" 감지된 오류는 {error_code}이며, 메시지는 '{error_message}'입니다."
        )

    if response_time is not None:
        summary_parts.append(
            f" 마지막 측정 응답 시간은 {float(response_time):.2f}초입니다."
        )

    return {
        "screen": st.session_state.get("last_screen", "화면 미확인"),
        "function": st.session_state.get("last_function", "기능 미확인"),
        "last_input": st.session_state.get("last_input", ""),
        "response_time": response_time,
        "error_code": error_code,
        "error_type": last_error.get("type", ""),
        "error_message": error_message,
        "action_history": "\n".join(action_lines) if action_lines else "기록된 최근 동작이 없습니다.",
        "auto_summary": "".join(summary_parts),
        "technical_traceback": last_error.get("traceback", ""),
    }


def save_error_report(report):
    ERROR_REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)

    with ERROR_REPORT_PATH.open("a", encoding="utf-8") as f:
        f.write(json.dumps(report, ensure_ascii=False) + "\n")

    reports = list(st.session_state.get("error_reports", []))
    reports.append(report)
    st.session_state["error_reports"] = reports[-100:]


@st.dialog("⚠️ 시스템 오류 자동 감지", width="large")
def show_error_report_dialog():
    draft = build_error_report_draft()

    st.warning("시스템이 오류 또는 이상 상태를 자동으로 감지했습니다.")
    st.caption(
        "오류 발생 당시의 화면, 기능, 입력값, 최근 동작, 에러 코드와 응답 시간을 자동으로 수집했습니다. "
        "내용을 확인하고 필요한 경우 추가 의견만 작성한 뒤 관리자에게 전송하세요."
    )

    c1, c2 = st.columns(2)

    with c1:
        st.text_input(
            "현재 화면",
            value=draft["screen"],
            disabled=True,
        )

    with c2:
        st.text_input(
            "사용 기능",
            value=draft["function"],
            disabled=True,
        )

    st.text_input(
        "최근 입력/검색어",
        value=draft["last_input"] or "기록 없음",
        disabled=True,
    )

    c3, c4 = st.columns(2)

    with c3:
        st.text_input(
            "에러 코드",
            value=draft["error_code"],
            disabled=True,
        )

    with c4:
        response_label = (
            f"{float(draft['response_time']):.2f}초"
            if draft["response_time"] is not None
            else "측정 기록 없음"
        )
        st.text_input(
            "최근 응답 시간",
            value=response_label,
            disabled=True,
        )

    st.text_area(
        "최근 동작 순서",
        value=draft["action_history"],
        height=160,
        disabled=True,
    )

    auto_summary = st.text_area(
        "자동 작성된 오류 내용",
        value=draft["auto_summary"],
        height=130,
        help="필요하면 사용자가 직접 수정할 수 있습니다.",
        key="error_report_auto_summary",
    )

    additional_note = st.text_area(
        "추가로 전달할 내용 (선택)",
        placeholder="예: 같은 버튼을 두 번 눌러도 동일한 오류가 발생했습니다.",
        key="error_report_additional_note",
    )

    st.checkbox(
        "신고 시점 화면 캡처 포함",
        value=False,
        disabled=True,
        help="화면 캡처는 2단계에서 브라우저 캡처 기능으로 연결할 예정입니다.",
    )

    st.caption(
        "🔒 주민등록번호 형태의 값은 자동 마스킹합니다. "
        "실제 사내 적용 시 추가 개인정보/기밀정보 마스킹 규칙을 확장할 수 있습니다."
    )

    if st.button(
        "📨 관리자에게 전송",
        type="primary",
        use_container_width=True,
        key="submit_error_report",
    ):
        report = {
            "report_id": "ERR-" + uuid.uuid4().hex[:8].upper(),
            "reported_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "screen": draft["screen"],
            "function": draft["function"],
            "last_input": draft["last_input"],
            "response_time": draft["response_time"],
            "error_code": draft["error_code"],
            "error_type": draft["error_type"],
            "error_message": draft["error_message"],
            "action_history": draft["action_history"],
            "auto_summary": _mask_sensitive_text(auto_summary),
            "additional_note": _mask_sensitive_text(additional_note),
            "screenshot": "2단계 연결 예정",
            "status": "미처리",
            "technical_traceback": draft["technical_traceback"],
        }

        try:
            save_error_report(report)
            st.session_state["pending_auto_error_dialog"] = False
            st.session_state["last_error"] = None
            st.success(
                f"✅ 오류 리포트가 관리자 오류함에 접수되었습니다. 신고번호: {report['report_id']}"
            )
        except Exception as e:
            st.error(f"신고 저장 중 오류가 발생했습니다: {e}")


@st.dialog("⚠️ 오류 신고", width="large")
def show_manual_error_report_dialog():
    """자동 감지되지 않은 이상 현상을 사용자가 직접 신고할 때 사용하는 팝업."""
    logs = st.session_state.get("error_action_logs", [])[-8:]
    action_lines = []

    for idx, event in enumerate(logs, start=1):
        line = (
            f"{idx}. [{event.get('time', '')}] "
            f"{event.get('screen', '')} · {event.get('action', '')}"
        )
        if event.get("input"):
            line += f" · 입력: {event.get('input', '')}"
        action_lines.append(line)

    current_screen = st.session_state.get("last_screen", "현재 화면")
    current_function = st.session_state.get("last_function", "앱 이용")
    last_input = st.session_state.get("last_input", "")
    last_response_time = st.session_state.get("last_response_time")

    st.info(
        "시스템이 자동으로 감지하지 못한 이상 현상이 있다면 직접 신고할 수 있습니다. "
        "최근 화면·기능·입력값·동작 기록은 자동으로 함께 첨부됩니다."
    )

    c1, c2 = st.columns(2)

    with c1:
        st.text_input(
            "현재 화면",
            value=current_screen,
            disabled=True,
            key="manual_report_screen",
        )

    with c2:
        st.text_input(
            "최근 사용 기능",
            value=current_function,
            disabled=True,
            key="manual_report_function",
        )

    st.text_input(
        "최근 입력/검색어",
        value=last_input or "기록 없음",
        disabled=True,
        key="manual_report_input",
    )

    if last_response_time is not None:
        st.caption(f"최근 측정 응답 시간: {float(last_response_time):.2f}초")

    if action_lines:
        st.text_area(
            "최근 동작 순서",
            value="\n".join(action_lines),
            height=150,
            disabled=True,
            key="manual_report_actions",
        )

    user_note = st.text_area(
        "어떤 문제가 있었나요?",
        placeholder="예: 질문하기 버튼을 눌렀는데 답변 내용이 이상하게 표시됐어요.",
        height=120,
        key="manual_error_report_note",
    )

    st.caption(
        "🔒 주민등록번호 형태의 값은 신고 저장 전에 자동 마스킹됩니다."
    )

    if st.button(
        "📨 관리자에게 신고",
        type="primary",
        use_container_width=True,
        key="manual_error_report_submit",
    ):
        if not user_note.strip():
            st.warning("신고할 내용을 간단히 입력해주세요.")
            return

        report = {
            "report_id": "ERR-" + uuid.uuid4().hex[:8].upper(),
            "reported_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "screen": current_screen,
            "function": current_function,
            "last_input": _mask_sensitive_text(last_input),
            "response_time": last_response_time,
            "error_code": "USER_REPORT",
            "error_type": "USER_REPORTED_ISSUE",
            "error_message": _mask_sensitive_text(user_note),
            "action_history": (
                "\n".join(action_lines)
                if action_lines
                else "기록된 최근 동작이 없습니다."
            ),
            "auto_summary": (
                f"{current_screen}에서 {current_function} 기능 사용 중 "
                "사용자가 이상 현상을 직접 신고했습니다."
            ),
            "additional_note": _mask_sensitive_text(user_note),
            "screenshot": "2단계 연결 예정",
            "status": "미처리",
            "technical_traceback": "",
        }

        try:
            save_error_report(report)
            st.success(
                f"✅ 신고가 접수되었습니다. 신고번호: {report['report_id']}"
            )
        except Exception as e:
            st.error(f"신고 저장 중 오류가 발생했습니다: {e}")


def render_floating_error_button(screen, function, key_suffix):
    """
    현재 탭 안에 플로팅 오류 신고 버튼을 표시합니다.
    버튼을 누르면 해당 탭의 화면/기능 정보가 신고서에 자동 반영됩니다.
    """
    st.markdown(
        """
        <style>
        [class*="st-key-floating_error_report_"] {
            position: fixed;
            right: 24px;
            bottom: 24px;
            z-index: 999999;
            width: auto !important;
        }

        [class*="st-key-floating_error_report_"] button {
            border-radius: 999px !important;
            padding: 0.65rem 1rem !important;
            box-shadow: 0 4px 18px rgba(0, 0, 0, 0.20);
            font-weight: 700;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    with st.container(key=f"floating_error_report_{key_suffix}"):
        if st.button(
            "⚠️ 오류 신고",
            key=f"floating_error_report_button_{key_suffix}",
            help="현재 화면의 오류 또는 이상 현상을 신고합니다.",
        ):
            # 신고 버튼이 눌린 '현재 탭'을 정확히 기록
            st.session_state["last_screen"] = screen
            st.session_state["last_function"] = function
            show_manual_error_report_dialog()


init_error_reporting()


# =========================================================
# 2. 공통 유틸리티
# =========================================================
def default_rows(columns, n=3):
    return pd.DataFrame([{c: "" for c in columns} for _ in range(n)])


def clean_records(df):
    """빈 행을 제거하고 dict 리스트로 변환."""
    if df is None:
        return []

    df = df.fillna("")
    records = []

    for row in df.to_dict(orient="records"):
        normalized = {
            str(k): str(v).strip()
            for k, v in row.items()
        }

        if any(normalized.values()):
            records.append(normalized)

    return records


def safe_value(value):
    """NaN/None 등을 빈 문자열로 정리."""
    if value is None:
        return ""

    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass

    return str(value).strip()


def read_excel_smart(uploaded_file):
    """
    제목 행/공백 행이 위에 있어도 실제 헤더 행을 자동으로 찾아 읽습니다.
    시연용 Excel은 3행에 실제 컬럼명이 있으므로 자동 감지됩니다.
    """
    raw = pd.read_excel(
        uploaded_file,
        header=None
    )

    known_headers = {
        "일자",
        "업무",
        "프로젝트/현장",
        "현재 단계",
        "구분",
        "회사/부서",
        "시스템/자산",
        "유형",
    }

    header_index = 0

    for i in range(
        min(
            len(raw),
            10
        )
    ):
        values = {
            str(value).strip()
            for value in raw.iloc[i].tolist()
            if pd.notna(value)
        }

        hit_count = len(
            known_headers & values
        )

        if hit_count >= 2:
            header_index = i
            break

    uploaded_file.seek(0)

    df = pd.read_excel(
        uploaded_file,
        header=header_index
    )

    df = df.dropna(
        how="all"
    )

    df = df.dropna(
        axis=1,
        how="all"
    )

    df.columns = [
        str(column).strip()
        for column in df.columns
    ]

    df = df.reset_index(
        drop=True
    )

    # Q&A에서 근거 Excel 행 번호를 계산할 때 사용
    df.attrs[
        "excel_header_row"
    ] = header_index + 1

    return df


# =========================================================
# 2. DOCX 생성 함수
# =========================================================
def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""

    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text is not None else "")

    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(9)


def add_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(title)
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(18)


def add_section_heading(doc, title):
    p = doc.add_paragraph()

    run = p.add_run(title)
    run.bold = True
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(11)


def add_key_value_table(doc, pairs, cols=2):
    rows = (len(pairs) + cols - 1) // cols

    table = doc.add_table(
        rows=rows,
        cols=cols * 2
    )
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    idx = 0

    for r in range(rows):
        for c in range(cols):
            if idx < len(pairs):
                key, value = pairs[idx]

                set_cell_text(
                    table.cell(r, c * 2),
                    key,
                    bold=True
                )

                set_cell_text(
                    table.cell(r, c * 2 + 1),
                    value
                )

                idx += 1

            else:
                set_cell_text(
                    table.cell(r, c * 2),
                    ""
                )

                set_cell_text(
                    table.cell(r, c * 2 + 1),
                    ""
                )

    return table


def add_records_table(doc, records, columns):
    table = doc.add_table(
        rows=1,
        cols=len(columns)
    )

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(columns):
        set_cell_text(
            table.rows[0].cells[i],
            col,
            bold=True
        )

    if records:
        for record in records:
            cells = table.add_row().cells

            for i, col in enumerate(columns):
                set_cell_text(
                    cells[i],
                    record.get(col, "")
                )

    else:
        cells = table.add_row().cells

        for i in range(len(columns)):
            set_cell_text(
                cells[i],
                ""
            )

    return table


def add_detail_table(doc, details):
    table = doc.add_table(
        rows=len(details),
        cols=2
    )

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (key, value) in enumerate(details):
        set_cell_text(
            table.cell(i, 0),
            key,
            bold=True
        )

        set_cell_text(
            table.cell(i, 1),
            value
        )

    return table


def build_docx(data):
    doc = Document()
    style_document(doc)

    add_title(
        doc,
        "업무 인수인계서"
    )

    add_key_value_table(
        doc,
        [
            ("기관명", data["meta"]["기관명"]),
            ("부서명", data["meta"]["부서명"]),
            ("문서번호", data["meta"]["문서번호"]),
            ("보존기간", data["meta"]["보존기간"]),
        ]
    )

    add_section_heading(
        doc,
        "I. 기본 정보"
    )

    add_key_value_table(
        doc,
        [
            ("소속 부서", data["basic"]["소속 부서"]),
            ("직위 / 직책", data["basic"]["직위 / 직책"]),
            ("인계자 성명", data["basic"]["인계자 성명"]),
            ("인수자 성명", data["basic"]["인수자 성명"]),
            ("작성일", data["basic"]["작성일"]),
            (
                "인수인계 완료 예정일",
                data["basic"]["인수인계 완료 예정일"]
            ),
        ]
    )

    add_section_heading(
        doc,
        "II. 담당업무 개요"
    )

    add_records_table(
        doc,
        data["tasks"],
        [
            "주요 업무",
            "업무 목적",
            "업무 프로세스",
            "우선순위",
            "비고",
        ]
    )

    add_section_heading(
        doc,
        "III. 업무 상세"
    )

    for i, item in enumerate(
        data["details"],
        start=1
    ):
        p = doc.add_paragraph()

        run = p.add_run(
            f"{i}) {item.get('업무명', '업무 상세')}"
        )

        run.bold = True
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "맑은 고딕"
        )
        run.font.size = Pt(10)

        add_detail_table(
            doc,
            [
                ("업무 개요", item.get("업무 개요", "")),
                ("목적 / 성과지표", item.get("목적 / 성과지표", "")),
                (
                    "진행 중 프로젝트 현황",
                    item.get("진행 중 프로젝트 현황", "")
                ),
                ("정기 업무", item.get("정기 업무", "")),
                ("비정기 업무", item.get("비정기 업무", "")),
                ("주요 일정 / 마감", item.get("주요 일정 / 마감", "")),
                (
                    "관련 시스템 / 계정 / 권한",
                    item.get("관련 시스템 / 계정 / 권한", "")
                ),
                (
                    "관련 담당자 / 연락처",
                    item.get("관련 담당자 / 연락처", "")
                ),
                ("협업 부서", item.get("협업 부서", "")),
                (
                    "특이사항 / 주의사항",
                    item.get("특이사항 / 주의사항", "")
                ),
                (
                    "리스크 / 미해결 이슈",
                    item.get("리스크 / 미해결 이슈", "")
                ),
                (
                    "참고 파일 경로 / 문서 링크",
                    item.get("참고 파일 경로 / 문서 링크", "")
                ),
                (
                    "후임자 숙지 필요사항",
                    item.get("후임자 숙지 필요사항", "")
                ),
                (
                    "인수인계 완료 여부",
                    item.get("인수인계 완료 여부", "")
                ),
            ]
        )

    add_section_heading(
        doc,
        "IV. 주요 일정 및 대외 커뮤니케이션"
    )

    p = doc.add_paragraph()
    r = p.add_run("1. 긴급 업무 일정")
    r.bold = True

    add_records_table(
        doc,
        data["urgent_schedule"],
        [
            "일자",
            "내용",
            "대응 방법",
            "담당",
        ]
    )

    p = doc.add_paragraph()
    r = p.add_run("2. 향후 1개월 주요 일정")
    r.bold = True

    add_records_table(
        doc,
        data["monthly_schedule"],
        [
            "일자",
            "일정 내용",
            "비고",
        ]
    )

    add_detail_table(
        doc,
        [
            (
                "3. 대외 커뮤니케이션 유의사항",
                data["communication_note"]
            )
        ]
    )

    add_section_heading(
        doc,
        "V. 계정 · 권한 · 자산 인계"
    )

    add_records_table(
        doc,
        data["assets"],
        [
            "시스템 / 자산명",
            "유형",
            "권한 수준",
            "인계 방법",
            "상태",
            "비고",
        ]
    )

    add_section_heading(
        doc,
        "VI. 인수인계 체크리스트"
    )

    add_records_table(
        doc,
        data["checklist"],
        [
            "체크",
            "항목",
            "확인일",
            "확인자",
            "비고",
        ]
    )

    add_section_heading(
        doc,
        "VII. 서명"
    )

    add_key_value_table(
        doc,
        [
            (
                "인계자 성명",
                data["signatures"]["인계자 성명"]
            ),
            (
                "인계자 서명일",
                data["signatures"]["인계자 서명일"]
            ),
            (
                "인수자 성명",
                data["signatures"]["인수자 성명"]
            ),
            (
                "인수자 서명일",
                data["signatures"]["인수자 서명일"]
            ),
            (
                "확인자(팀장 등) 성명",
                data["signatures"]["확인자 성명"]
            ),
            (
                "확인자 서명일",
                data["signatures"]["확인자 서명일"]
            ),
        ]
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(
        "본 문서는 업무 연속성을 위해 작성된 내부 인수인계 자료입니다. "
        "외부 반출 시 소속 부서의 사전 승인이 필요합니다."
    )

    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "맑은 고딕"
    )
    run.font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


# =========================================================
# 3. 인수인계 준비도 계산
# =========================================================
def calculate_readiness(reviewed_data):
    missing_items = []
    check_items = []

    def check_value(label, value):
        if (
            value is not None
            and str(value).strip() != ""
        ):
            check_items.append(True)

        else:
            check_items.append(False)
            missing_items.append(label)

    # 기본 정보
    check_value(
        "소속 부서",
        reviewed_data["basic"].get(
            "소속 부서",
            ""
        )
    )

    check_value(
        "직위 / 직책",
        reviewed_data["basic"].get(
            "직위 / 직책",
            ""
        )
    )

    check_value(
        "인계자 성명",
        reviewed_data["basic"].get(
            "인계자 성명",
            ""
        )
    )

    check_value(
        "인수자 성명",
        reviewed_data["basic"].get(
            "인수자 성명",
            ""
        )
    )

    check_value(
        "작성일",
        reviewed_data["basic"].get(
            "작성일",
            ""
        )
    )

    check_value(
        "인수인계 완료 예정일",
        reviewed_data["basic"].get(
            "인수인계 완료 예정일",
            ""
        )
    )

    # 담당업무 개요
    if len(reviewed_data["tasks"]) > 0:
        check_items.append(True)

    else:
        check_items.append(False)
        missing_items.append(
            "담당업무 개요"
        )

    # 업무 상세
    if len(reviewed_data["details"]) == 0:
        missing_items.append(
            "업무 상세"
        )
        check_items.extend(
            [False] * 7
        )

    else:
        for i, detail in enumerate(
            reviewed_data["details"],
            start=1
        ):
            check_value(
                f"업무 {i} - 업무명",
                detail.get(
                    "업무명",
                    ""
                )
            )

            check_value(
                f"업무 {i} - 업무 개요",
                detail.get(
                    "업무 개요",
                    ""
                )
            )

            check_value(
                f"업무 {i} - 주요 일정 / 마감",
                detail.get(
                    "주요 일정 / 마감",
                    ""
                )
            )

            check_value(
                f"업무 {i} - 관련 담당자 / 연락처",
                detail.get(
                    "관련 담당자 / 연락처",
                    ""
                )
            )

            check_value(
                f"업무 {i} - 리스크 / 미해결 이슈",
                detail.get(
                    "리스크 / 미해결 이슈",
                    ""
                )
            )

            check_value(
                f"업무 {i} - 참고 파일",
                detail.get(
                    "참고 파일 경로 / 문서 링크",
                    ""
                )
            )

            check_value(
                f"업무 {i} - 후임자 숙지사항",
                detail.get(
                    "후임자 숙지 필요사항",
                    ""
                )
            )

    # 일정
    if len(
        reviewed_data["urgent_schedule"]
    ) > 0:
        check_items.append(True)

    else:
        check_items.append(False)
        missing_items.append(
            "긴급 업무 일정"
        )

    if len(
        reviewed_data["monthly_schedule"]
    ) > 0:
        check_items.append(True)

    else:
        check_items.append(False)
        missing_items.append(
            "향후 1개월 주요 일정"
        )

    # 자산
    if len(
        reviewed_data["assets"]
    ) > 0:
        check_items.append(True)

    else:
        check_items.append(False)
        missing_items.append(
            "계정 · 권한 · 자산 인계"
        )

    total_count = len(check_items)
    completed_count = sum(check_items)

    if total_count > 0:
        readiness = int(
            completed_count
            / total_count
            * 100
        )

    else:
        readiness = 0

    return (
        readiness,
        completed_count,
        missing_items
    )


def show_readiness(reviewed_data):
    readiness, completed_count, missing_items = (
        calculate_readiness(reviewed_data)
    )

    st.subheader("📊 인수인계 준비도")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric(
            "인수인계 준비도",
            f"{readiness}%"
        )

    with col2:
        st.metric(
            "작성 완료 항목",
            f"{completed_count}개"
        )

    with col3:
        st.metric(
            "보완 필요 항목",
            f"{len(missing_items)}개"
        )

    st.progress(
        readiness / 100
    )

    if readiness >= 90:
        st.success(
            "✅ 인수인계 준비 상태가 매우 좋습니다."
        )

    elif readiness >= 70:
        st.warning(
            "⚠️ 일부 항목의 보완이 필요합니다."
        )

    else:
        st.error(
            "🚨 인수인계에 필요한 정보가 부족합니다."
        )

    if missing_items:
        with st.expander(
            "⚠️ 보완이 필요한 항목 보기",
            expanded=True
        ):
            for item in missing_items:
                st.write(
                    f"• {item}"
                )

    else:
        st.success(
            "🎉 필수 인수인계 항목이 모두 작성되었습니다!"
        )


# =========================================================
# 4. Excel 업무자료 자동 분류 함수
# =========================================================
def classify_excel_files(excel_data):
    schedule_df = None
    project_df = None
    contact_df = None
    asset_df = None
    unknown_files = []

    for file_name, df in excel_data.items():
        normalized_name = file_name.replace(
            " ",
            ""
        )

        if "업무일정" in normalized_name:
            schedule_df = df

        elif "프로젝트" in normalized_name:
            project_df = df

        elif "담당자" in normalized_name:
            contact_df = df

        elif (
            "계정" in normalized_name
            or "권한" in normalized_name
            or "자산" in normalized_name
        ):
            asset_df = df

        else:
            unknown_files.append(
                file_name
            )

    return (
        schedule_df,
        project_df,
        contact_df,
        asset_df,
        unknown_files,
    )


# =========================================================
# 5. 후임자 대시보드
# =========================================================
def _to_number(value):
    try:
        return float(value)
    except Exception:
        return 0.0


def _to_datetime(value):
    if value is None:
        return pd.NaT

    if isinstance(value, pd.Timestamp):
        return value

    try:
        if pd.isna(value):
            return pd.NaT
    except Exception:
        pass

    # Excel serial date 대응
    if isinstance(value, (int, float)):
        if 30000 <= float(value) <= 70000:
            try:
                return pd.to_datetime(
                    float(value),
                    unit="D",
                    origin="1899-12-30"
                )
            except Exception:
                return pd.NaT

    return pd.to_datetime(
        value,
        errors="coerce"
    )


def _format_dashboard_date(value):
    dt = _to_datetime(value)

    if pd.isna(dt):
        return safe_value(value)

    return dt.strftime("%Y-%m-%d")


def build_dashboard_summary(
    schedule_df,
    project_df,
    asset_df
):
    """Excel 자료를 이용해 후임자용 핵심 지표를 계산."""
    summary = {
        "urgent_count": 0,
        "open_work_count": 0,
        "issue_count": 0,
        "pending_asset_count": 0,
        "average_progress": 0,
        "nearest_deadline": "",
        "nearest_work": "",
        "nearest_project": "",
    }

    if schedule_df is not None and not schedule_df.empty:
        if "우선순위" in schedule_df.columns:
            summary["urgent_count"] = int(
                schedule_df["우선순위"]
                .astype(str)
                .str.strip()
                .eq("긴급")
                .sum()
            )

        if "상태" in schedule_df.columns:
            completed_labels = {
                "완료",
                "종료",
                "제출완료",
            }

            summary["open_work_count"] = int(
                (
                    ~schedule_df["상태"]
                    .astype(str)
                    .str.strip()
                    .isin(completed_labels)
                ).sum()
            )

        if "진행률" in schedule_df.columns:
            progress = pd.to_numeric(
                schedule_df["진행률"],
                errors="coerce"
            )

            if progress.notna().any():
                summary["average_progress"] = int(
                    round(
                        progress.mean()
                    )
                )

        if "마감" in schedule_df.columns:
            deadlines = schedule_df.copy()

            deadlines["_deadline"] = deadlines[
                "마감"
            ].apply(
                _to_datetime
            )

            deadlines = deadlines[
                deadlines["_deadline"].notna()
            ]

            if not deadlines.empty:
                nearest = deadlines.sort_values(
                    "_deadline"
                ).iloc[0]

                summary["nearest_deadline"] = (
                    _format_dashboard_date(
                        nearest.get(
                            "마감",
                            ""
                        )
                    )
                )

                summary["nearest_work"] = safe_value(
                    nearest.get(
                        "업무",
                        ""
                    )
                )

                summary["nearest_project"] = safe_value(
                    nearest.get(
                        "프로젝트/현장",
                        ""
                    )
                )

    if project_df is not None and not project_df.empty:
        if "리스크/이슈" in project_df.columns:
            issue_series = (
                project_df["리스크/이슈"]
                .fillna("")
                .astype(str)
                .str.strip()
            )

            summary["issue_count"] = int(
                issue_series.ne("").sum()
            )

    if asset_df is not None and not asset_df.empty:
        if "상태" in asset_df.columns:
            summary["pending_asset_count"] = int(
                (
                    ~asset_df["상태"]
                    .astype(str)
                    .str.strip()
                    .eq("완료")
                ).sum()
            )

    return summary


def show_handover_dashboard(
    schedule_df,
    project_df,
    contact_df,
    asset_df
):
    uploaded_count = sum(
        df is not None
        for df in [
            schedule_df,
            project_df,
            contact_df,
            asset_df,
        ]
    )

    st.header(
        "🏠 후임자 업무 대시보드"
    )

    st.caption(
        "후임자가 인수 직후 확인해야 할 업무, 마감, 리스크와 "
        "권한 인계 상태를 한 화면에서 확인합니다."
    )

    if uploaded_count < 4:
        st.info(
            "먼저 **📂 Excel 업무자료** 탭에서 "
            "01~04 Excel 파일을 모두 업로드해주세요. "
            "업로드가 완료되면 이 대시보드가 자동으로 채워집니다."
        )

        c1, c2, c3, c4 = st.columns(4)

        labels = [
            ("업무일정", schedule_df),
            ("프로젝트", project_df),
            ("담당자", contact_df),
            ("권한·자산", asset_df),
        ]

        for column, (
            label,
            df
        ) in zip(
            [c1, c2, c3, c4],
            labels
        ):
            with column:
                st.metric(
                    label,
                    "✅ 연결"
                    if df is not None
                    else "➖ 대기"
                )

        return

    summary = build_dashboard_summary(
        schedule_df,
        project_df,
        asset_df
    )

    # 핵심 KPI
    c1, c2, c3, c4, c5 = st.columns(5)

    with c1:
        st.metric(
            "🚨 긴급 업무",
            f"{summary['urgent_count']}건"
        )

    with c2:
        st.metric(
            "🟡 진행/예정 업무",
            f"{summary['open_work_count']}건"
        )

    with c3:
        st.metric(
            "⚠️ 미해결 이슈",
            f"{summary['issue_count']}건"
        )

    with c4:
        st.metric(
            "🔐 미완료 인계",
            f"{summary['pending_asset_count']}건"
        )

    with c5:
        st.metric(
            "📈 평균 진행률",
            f"{summary['average_progress']}%"
        )

    st.divider()

    # 가장 가까운 마감
    st.subheader(
        "⏰ 가장 가까운 마감"
    )

    if summary["nearest_deadline"]:
        st.warning(
            f"**{summary['nearest_deadline']}**까지 "
            f"**{summary['nearest_project']} - "
            f"{summary['nearest_work']}** 업무를 확인해야 합니다."
        )

    else:
        st.info(
            "등록된 마감 일정이 없습니다."
        )

    # 다음 우선 업무 3건
    st.subheader(
        "📌 후임자가 먼저 확인할 업무"
    )

    schedule_view = schedule_df.copy()

    schedule_view["_deadline"] = schedule_view[
        "마감"
    ].apply(
        _to_datetime
    )

    priority_rank = {
        "긴급": 4,
        "상": 3,
        "중": 2,
        "하": 1,
    }

    schedule_view["_priority"] = (
        schedule_view["우선순위"]
        .astype(str)
        .str.strip()
        .map(priority_rank)
        .fillna(0)
    )

    schedule_view = schedule_view.sort_values(
        [
            "_deadline",
            "_priority",
        ],
        ascending=[
            True,
            False,
        ]
    ).head(3)

    priority_rows = []

    for _, row in schedule_view.iterrows():
        priority_rows.append(
            {
                "마감": _format_dashboard_date(
                    row.get(
                        "마감",
                        ""
                    )
                ),
                "업무": safe_value(
                    row.get(
                        "업무",
                        ""
                    )
                ),
                "프로젝트/현장": safe_value(
                    row.get(
                        "프로젝트/현장",
                        ""
                    )
                ),
                "우선순위": safe_value(
                    row.get(
                        "우선순위",
                        ""
                    )
                ),
                "진행률": (
                    f"{int(_to_number(row.get('진행률', 0)))}%"
                ),
                "다음 조치": safe_value(
                    row.get(
                        "다음 조치",
                        ""
                    )
                ),
            }
        )

    st.dataframe(
        pd.DataFrame(
            priority_rows
        ),
        use_container_width=True,
        hide_index=True
    )

    # 프로젝트 상태
    st.subheader(
        "📊 프로젝트 진행 현황"
    )

    project_columns = [
        column
        for column in [
            "프로젝트/현장",
            "현재 단계",
            "진행률",
            "다음 액션",
            "리스크/이슈",
            "마감일",
        ]
        if column in project_df.columns
    ]

    project_view = project_df[
        project_columns
    ].copy()

    if "진행률" in project_view.columns:
        project_view[
            "진행률"
        ] = project_view[
            "진행률"
        ].apply(
            lambda value: (
                f"{int(_to_number(value))}%"
            )
        )

    if "마감일" in project_view.columns:
        project_view[
            "마감일"
        ] = project_view[
            "마감일"
        ].apply(
            _format_dashboard_date
        )

    st.dataframe(
        project_view,
        use_container_width=True,
        hide_index=True
    )

    left, right = st.columns(2)

    # 미해결 이슈
    with left:
        st.subheader(
            "⚠️ 미해결 이슈"
        )

        issues = []

        for _, row in project_df.iterrows():
            issue = safe_value(
                row.get(
                    "리스크/이슈",
                    ""
                )
            )

            if issue:
                issues.append(
                    (
                        safe_value(
                            row.get(
                                "프로젝트/현장",
                                ""
                            )
                        ),
                        issue,
                    )
                )

        if issues:
            for project, issue in issues:
                st.write(
                    f"• **{project}** — {issue}"
                )

        else:
            st.success(
                "등록된 미해결 이슈가 없습니다."
            )

    # 미완료 권한/자산
    with right:
        st.subheader(
            "🔐 남아 있는 인계"
        )

        pending_assets = asset_df[
            ~asset_df["상태"]
            .astype(str)
            .str.strip()
            .eq("완료")
        ]

        if not pending_assets.empty:
            for _, row in pending_assets.iterrows():
                st.write(
                    f"• **{safe_value(row.get('시스템/자산', ''))}** "
                    f"— {safe_value(row.get('상태', ''))} "
                    f"/ 목표 {_format_dashboard_date(row.get('완료 목표일', ''))}"
                )

        else:
            st.success(
                "계정·권한·자산 인계가 모두 완료되었습니다."
            )

    st.divider()

    st.success(
        "💡 위 내용을 확인한 뒤 **💬 후임자 Q&A** 탭에서 "
        "세부 업무를 질문하거나, **🤖 업무메모 자동 인수인계**에서 "
        "최종 인수인계서를 생성할 수 있습니다."
    )


# =========================================================
# 5. 메인 화면 탭
# =========================================================
schedule_df = None
project_df = None
contact_df = None
asset_df = None

(
    tab_dashboard,
    tab_excel,
    tab_qa,
    tab_onboarding_schedule,
    tab_memo,
    tab_manual,
    tab_error_admin,
) = st.tabs(
    [
        "🏠 대시보드",
        "📂 Excel 업무자료",
        "💬 후임자 Q&A",
        "📅 온보딩 일과표",
        "🤖 업무메모 자동 인수인계",
        "✍️ 직접 작성",
        "🛠 관리자 오류함",
    ]
)


# =========================================================
# TAB 1. Excel 업무자료 업로드
# =========================================================
with tab_excel:
    render_floating_error_button(
        screen="📂 Excel 업무자료",
        function="업무자료 업로드 및 자동 분류",
        key_suffix="excel",
    )
    st.header("📂 업무자료 통합 업로드")

    st.info(
        "업무일정, 프로젝트 진행현황, 담당자 연락망, "
        "계정·권한·자산 Excel 파일을 한 번에 업로드할 수 있습니다."
    )

    uploaded_excels = st.file_uploader(
        "Excel 업무자료를 업로드하세요.",
        type=["xlsx"],
        accept_multiple_files=True,
        key="excel_upload"
    )

    excel_data = {}

    if uploaded_excels:
        uploaded_names = [f.name for f in uploaded_excels]

        if uploaded_names != st.session_state.get("last_uploaded_files", []):
            track_action(
                screen="📂 Excel 업무자료",
                function="업무자료 업로드",
                action="Excel 파일 업로드",
                input_value=", ".join(uploaded_names),
                metadata={"file_count": len(uploaded_names)},
            )
            st.session_state["last_uploaded_files"] = uploaded_names

        for uploaded_file in uploaded_excels:
            try:
                df = read_excel_smart(
                    uploaded_file
                )

                excel_data[
                    uploaded_file.name
                ] = df

            except Exception as e:
                record_error(
                    screen="📂 Excel 업무자료",
                    function="Excel 파일 읽기",
                    error=e,
                    code="EXCEL_READ_ERROR",
                    input_value=uploaded_file.name,
                )
                st.error(
                    f"❌ {uploaded_file.name} 읽기 실패 · 시스템이 오류를 자동 감지했습니다."
                )

        (
            schedule_df,
            project_df,
            contact_df,
            asset_df,
            unknown_files,
        ) = classify_excel_files(
            excel_data
        )

        st.subheader("✅ 자료 자동 분류 결과")

        c1, c2, c3, c4 = st.columns(4)

        with c1:
            st.metric(
                "업무일정",
                "업로드 완료"
                if schedule_df is not None
                else "미업로드"
            )

        with c2:
            st.metric(
                "프로젝트 현황",
                "업로드 완료"
                if project_df is not None
                else "미업로드"
            )

        with c3:
            st.metric(
                "담당자 연락망",
                "업로드 완료"
                if contact_df is not None
                else "미업로드"
            )

        with c4:
            st.metric(
                "계정·권한·자산",
                "업로드 완료"
                if asset_df is not None
                else "미업로드"
            )

        if unknown_files:
            st.warning(
                "자동 분류하지 못한 파일: "
                + ", ".join(
                    unknown_files
                )
            )

        st.divider()
        st.subheader("📋 업로드된 업무자료 미리보기")

        for file_name, df in excel_data.items():
            with st.expander(
                f"📄 {file_name}",
                expanded=False
            ):
                st.caption(
                    f"{len(df)}행 × {len(df.columns)}열"
                )

                st.dataframe(
                    df,
                    use_container_width=True
                )

        st.success(
            "Excel 자료 업로드 및 자동 분류까지 완료되었습니다. "
            "다음 단계에서 이 자료를 기반으로 후임자 Q&A 기능을 연결할 수 있습니다."
        )

    else:
        st.caption(
            "권장 파일: 01_업무일정.xlsx / "
            "02_프로젝트_진행현황.xlsx / "
            "03_담당자_연락망.xlsx / "
            "04_계정_권한_자산.xlsx"
        )


# =========================================================
# TAB 2. 후임자 대시보드
# =========================================================
with tab_dashboard:
    render_floating_error_button(
        screen="🏠 대시보드",
        function="후임자 업무 대시보드",
        key_suffix="dashboard",
    )
    show_handover_dashboard(
        schedule_df,
        project_df,
        contact_df,
        asset_df
    )


# =========================================================
# TAB 2. 후임자 Q&A
# =========================================================
with tab_qa:
    render_floating_error_button(
        screen="💬 후임자 Q&A",
        function="후임자 질문 검색",
        key_suffix="qa",
    )
    st.header("💬 후임자 Q&A")

    st.caption(
        "업로드된 업무일정, 프로젝트 현황, 담당자 연락망, "
        "계정·권한·자산 자료를 검색해 답변합니다. "
        "현재 버전은 OpenAI API 없이 규칙 기반으로 동작합니다."
    )

    uploaded_count = sum(
        df is not None
        for df in [
            schedule_df,
            project_df,
            contact_df,
            asset_df,
        ]
    )

    if uploaded_count < 4:
        st.warning(
            "먼저 📂 Excel 업무자료 탭에서 "
            "01~04 Excel 파일을 모두 업로드해주세요."
        )

    else:
        c1, c2, c3, c4 = st.columns(4)

        c1.success("✅ 업무일정")
        c2.success("✅ 프로젝트")
        c3.success("✅ 담당자")
        c4.success("✅ 권한·자산")

        st.markdown("#### 질문 예시")

        sample_questions = [
            "직접 입력",
            "9월 1일에 가장 먼저 해야 할 업무가 뭐야?",
            "A동 보고서는 지금 어디까지 진행됐어?",
            "A동 보고서가 늦어질 수 있는 이유가 있어?",
            "B공장 견적이 아직 완료되지 않은 이유는 뭐야?",
            "B공장 견적 다음 액션은?",
            "C센터 회의 전에 뭘 준비해야 해?",
            "C센터에서 아직 결정 안 된 게 뭐야?",
            "A동 고객사 담당자는 누구야?",
            "박현우 부장님께 보고서 보낼 때 주의할 점 있어?",
            "감지기 단가 문의는 누구한테 해야 해?",
            "후임자가 공용드라이브를 바로 쓸 수 있어?",
            "비밀번호를 인계자한테 받아야 하는 시스템이 있어?",
            "아직 완료되지 않은 권한 인계 알려줘",
            "진행률 50% 이상인 프로젝트 알려줘",
            "후임자가 가장 먼저 확인할 3가지만 뽑아줘",
        ]

        selected_question = st.selectbox(
            "예시 질문 선택",
            sample_questions,
            key="qa_sample_question"
        )

        typed_question = st.text_input(
            "후임자가 궁금한 내용을 입력하세요",
            placeholder="예: A동 보고서는 지금 어디까지 진행됐어?",
            key="qa_question"
        )

        if typed_question.strip():
            final_question = typed_question.strip()

        elif selected_question != "직접 입력":
            final_question = selected_question

        else:
            final_question = ""

        ask_button = st.button(
            "🔎 질문하기",
            type="primary",
            use_container_width=True,
            key="qa_ask_button"
        )

        if ask_button:
            if not final_question:
                st.warning(
                    "질문을 입력하거나 예시 질문을 선택해주세요."
                )

            else:
                with st.chat_message("user"):
                    st.write(
                        final_question
                    )

                track_action(
                    screen="💬 후임자 Q&A",
                    function="후임자 질문 검색",
                    action="질문하기 버튼 클릭",
                    input_value=final_question,
                )

                qa_started_at = time.perf_counter()
                result = None

                try:
                    result = answer_question(
                        final_question,
                        schedule_df=schedule_df,
                        project_df=project_df,
                        contact_df=contact_df,
                        asset_df=asset_df,
                    )

                    qa_elapsed = time.perf_counter() - qa_started_at
                    st.session_state["last_response_time"] = qa_elapsed

                    track_action(
                        screen="💬 후임자 Q&A",
                        function="후임자 질문 검색",
                        action="답변 생성 완료",
                        input_value=final_question,
                        metadata={"response_time": round(qa_elapsed, 3)},
                    )

                    if qa_elapsed >= SLOW_RESPONSE_SECONDS:
                        record_system_issue(
                            screen="💬 후임자 Q&A",
                            function="후임자 질문 검색",
                            code="QA_SLOW_RESPONSE",
                            message=(
                                f"Q&A 응답 시간이 기준({SLOW_RESPONSE_SECONDS:.0f}초)을 "
                                f"초과했습니다: {qa_elapsed:.2f}초"
                            ),
                            issue_type="SLOW_RESPONSE",
                            input_value=final_question,
                            response_time=qa_elapsed,
                        )

                except Exception as e:
                    qa_elapsed = time.perf_counter() - qa_started_at

                    record_error(
                        screen="💬 후임자 Q&A",
                        function="후임자 질문 검색",
                        error=e,
                        code="QA_ANSWER_ERROR",
                        input_value=final_question,
                        response_time=qa_elapsed,
                    )

                    st.error(
                        "❌ 답변 생성 중 오류가 발생했습니다. "
                        "시스템이 자동 감지했으며 오류 리포트 창이 열립니다."
                    )

                if result is not None:
                    with st.chat_message("assistant"):
                        st.markdown(
                            result["answer"]
                        )

                        if result["sources"]:
                            st.markdown(
                                "##### 📎 답변 근거"
                            )

                            for source in result["sources"]:
                                location = (
                                    f"{source['sheet']} "
                                    f"{source['row']}행"
                                )

                                st.caption(
                                    f"• {source['file']} · {location}"
                                )

                                if source.get(
                                    "evidence"
                                ):
                                    st.caption(
                                        f"  ↳ {source['evidence']}"
                                    )

                        else:
                            st.caption(
                                "연결된 근거 자료가 없습니다."
                            )

        st.divider()

        with st.expander(
            "ℹ️ 현재 Q&A 방식"
        ):
            st.write(
                "현재 버전은 질문의 키워드와 현장명/담당자명 등을 분석해 "
                "관련 Excel 행을 찾고, 정해진 답변 템플릿으로 응답합니다."
            )

            st.write(
                "향후 LLM API 또는 RAG를 연결하면 표현이 달라도 문맥을 이해하고 "
                "여러 문서를 종합하는 자유로운 질의응답으로 확장할 수 있습니다."
            )


# =========================================================
# TAB 2. 업무메모 자동 인수인계

# =========================================================
# TAB. 후임자 온보딩 일과표 자동 생성기 (팀원 기능 통합)
# =========================================================
with tab_onboarding_schedule:
    render_floating_error_button(
        screen="📅 온보딩 일과표",
        function="온보딩 일과표 생성 및 수정",
        key_suffix="onboarding",
    )
    st.header("📅 후임자 온보딩 일과표")
    st.caption(
        "팀원이 제작한 일정표 자동 생성 모듈을 Streamlit 안에 통합했습니다. "
        "회의록·캘린더·메일 DOCX를 업로드하면 1주차와 2주차 온보딩 일과표를 자동 생성합니다."
    )

    st.info(
        "💡 **사용 방법:** 아래 화면에서 회의록(.docx) 여러 개, "
        "캘린더 문서(.docx) 1개, 메일(.docx) 여러 개를 넣고 "
        "**일정표 생성**을 누르세요. 생성된 일정의 업무명은 화면에서 직접 수정할 수 있습니다."
    )

    components.html(
        ONBOARDING_SCHEDULE_HTML,
        height=1450,
        scrolling=True,
    )


# =========================================================
with tab_memo:
    render_floating_error_button(
        screen="🤖 업무메모 자동 인수인계",
        function="업무메모 자동 인수인계",
        key_suffix="memo",
    )
    st.header("🤖 업무메모 자동 인수인계")

    uploaded_memo = st.file_uploader(
        "업무메모.txt 파일을 업로드하세요",
        type=["txt"],
        key="auto_memo_upload"
    )

    if uploaded_memo is not None:
        try:
            # 1. 메모 읽기
            memo_text = uploaded_memo.getvalue().decode(
                "utf-8-sig"
            )

            # 2. 규칙 기반 분석
            parsed_data = parse_memo_text(
                memo_text
            )

            st.success(
                "✅ 업무메모 분석 완료! 아래 내용을 확인하고 수정해주세요."
            )

            with st.expander(
                "📋 원본 업무메모 보기"
            ):
                st.text(
                    memo_text
                )

            # -----------------------------------------
            # I. 기본 정보
            # -----------------------------------------
            st.subheader(
                "Ⅰ. 기본 정보"
            )

            c1, c2 = st.columns(2)

            with c1:
                auto_institution = st.text_input(
                    "기관명",
                    value=parsed_data[
                        "meta"
                    ].get(
                        "기관명",
                        ""
                    ),
                    key="auto_institution"
                )

                auto_department_meta = st.text_input(
                    "부서명",
                    value=parsed_data[
                        "meta"
                    ].get(
                        "부서명",
                        ""
                    ),
                    key="auto_department_meta"
                )

                auto_department = st.text_input(
                    "소속 부서",
                    value=parsed_data[
                        "basic"
                    ].get(
                        "소속 부서",
                        ""
                    ),
                    key="auto_department"
                )

                auto_giver = st.text_input(
                    "인계자 성명",
                    value=parsed_data[
                        "basic"
                    ].get(
                        "인계자 성명",
                        ""
                    ),
                    key="auto_giver"
                )

                auto_written_date = st.text_input(
                    "작성일",
                    value=parsed_data[
                        "basic"
                    ].get(
                        "작성일",
                        ""
                    ),
                    key="auto_written_date"
                )

            with c2:
                auto_document_no = st.text_input(
                    "문서번호",
                    value=parsed_data[
                        "meta"
                    ].get(
                        "문서번호",
                        ""
                    ),
                    key="auto_document_no"
                )

                auto_retention = st.text_input(
                    "보존기간",
                    value=parsed_data[
                        "meta"
                    ].get(
                        "보존기간",
                        ""
                    ),
                    key="auto_retention"
                )

                auto_position = st.text_input(
                    "직위 / 직책",
                    value=parsed_data[
                        "basic"
                    ].get(
                        "직위 / 직책",
                        ""
                    ),
                    key="auto_position"
                )

                auto_receiver = st.text_input(
                    "인수자 성명",
                    value=parsed_data[
                        "basic"
                    ].get(
                        "인수자 성명",
                        ""
                    ),
                    key="auto_receiver"
                )

                auto_expected_date = st.text_input(
                    "인수인계 완료 예정일",
                    value=parsed_data[
                        "basic"
                    ].get(
                        "인수인계 완료 예정일",
                        ""
                    ),
                    key="auto_expected_date"
                )

            # -----------------------------------------
            # II. 담당업무 개요
            # -----------------------------------------
            st.subheader(
                "Ⅱ. 담당업무 개요"
            )

            task_columns = [
                "주요 업무",
                "업무 목적",
                "업무 프로세스",
                "우선순위",
                "비고",
            ]

            auto_tasks_df = st.data_editor(
                pd.DataFrame(
                    parsed_data.get(
                        "tasks",
                        []
                    ),
                    columns=task_columns
                ),
                num_rows="dynamic",
                use_container_width=True,
                key="auto_tasks"
            )

            # -----------------------------------------
            # III. 업무 상세
            # -----------------------------------------
            st.subheader(
                "Ⅲ. 업무 상세"
            )

            edited_details = []

            detail_fields = [
                "업무 개요",
                "목적 / 성과지표",
                "진행 중 프로젝트 현황",
                "정기 업무",
                "비정기 업무",
                "주요 일정 / 마감",
                "관련 시스템 / 계정 / 권한",
                "관련 담당자 / 연락처",
                "협업 부서",
                "특이사항 / 주의사항",
                "리스크 / 미해결 이슈",
                "참고 파일 경로 / 문서 링크",
                "후임자 숙지 필요사항",
            ]

            for i, detail in enumerate(
                parsed_data.get(
                    "details",
                    []
                )
            ):
                task_name = detail.get(
                    "업무명",
                    f"업무 {i + 1}"
                )

                with st.expander(
                    f"📌 {i + 1}. {task_name}",
                    expanded=True
                ):
                    edited_task_name = st.text_input(
                        "업무명",
                        value=task_name,
                        key=f"auto_detail_name_{i}"
                    )

                    edited_detail = {
                        "업무명": edited_task_name
                    }

                    for field in detail_fields:
                        edited_detail[
                            field
                        ] = st.text_area(
                            field,
                            value=detail.get(
                                field,
                                ""
                            ),
                            key=f"auto_detail_{i}_{field}"
                        )

                    status_options = [
                        "미완료",
                        "진행중",
                        "완료",
                    ]

                    current_status = detail.get(
                        "인수인계 완료 여부",
                        "미완료"
                    )

                    if current_status not in status_options:
                        current_status = "미완료"

                    edited_detail[
                        "인수인계 완료 여부"
                    ] = st.selectbox(
                        "인수인계 완료 여부",
                        status_options,
                        index=status_options.index(
                            current_status
                        ),
                        key=f"auto_done_{i}"
                    )

                    edited_details.append(
                        edited_detail
                    )

            # -----------------------------------------
            # IV. 일정 및 커뮤니케이션
            # -----------------------------------------
            st.subheader(
                "Ⅳ. 주요 일정 및 대외 커뮤니케이션"
            )

            st.markdown(
                "#### 1. 긴급 업무 일정"
            )

            urgent_columns = [
                "일자",
                "내용",
                "대응 방법",
                "담당",
            ]

            auto_urgent_df = st.data_editor(
                pd.DataFrame(
                    parsed_data.get(
                        "urgent_schedule",
                        []
                    ),
                    columns=urgent_columns
                ),
                num_rows="dynamic",
                use_container_width=True,
                key="auto_urgent"
            )

            st.markdown(
                "#### 2. 향후 1개월 주요 일정"
            )

            monthly_columns = [
                "일자",
                "일정 내용",
                "비고",
            ]

            auto_monthly_df = st.data_editor(
                pd.DataFrame(
                    parsed_data.get(
                        "monthly_schedule",
                        []
                    ),
                    columns=monthly_columns
                ),
                num_rows="dynamic",
                use_container_width=True,
                key="auto_monthly"
            )

            st.markdown(
                "#### 3. 대외 커뮤니케이션 유의사항"
            )

            auto_communication = st.text_area(
                "커뮤니케이션 유의사항",
                value=parsed_data.get(
                    "communication_note",
                    ""
                ),
                key="auto_communication"
            )

            # -----------------------------------------
            # V. 자산
            # -----------------------------------------
            st.subheader(
                "Ⅴ. 계정 · 권한 · 자산 인계"
            )

            asset_columns = [
                "시스템 / 자산명",
                "유형",
                "권한 수준",
                "인계 방법",
                "상태",
                "비고",
            ]

            auto_assets_df = st.data_editor(
                pd.DataFrame(
                    parsed_data.get(
                        "assets",
                        []
                    ),
                    columns=asset_columns
                ),
                num_rows="dynamic",
                use_container_width=True,
                key="auto_assets"
            )

            # -----------------------------------------
            # VI. 체크리스트
            # -----------------------------------------
            st.subheader(
                "Ⅵ. 인수인계 체크리스트"
            )

            checklist_columns = [
                "체크",
                "항목",
                "확인일",
                "확인자",
                "비고",
            ]

            auto_checklist_df = st.data_editor(
                pd.DataFrame(
                    parsed_data.get(
                        "checklist",
                        []
                    ),
                    columns=checklist_columns
                ),
                num_rows="dynamic",
                use_container_width=True,
                key="auto_checklist"
            )

            # -----------------------------------------
            # VII. 서명
            # -----------------------------------------
            st.subheader(
                "Ⅶ. 서명"
            )

            s1, s2, s3 = st.columns(3)

            with s1:
                auto_sign_giver = st.text_input(
                    "인계자",
                    value=parsed_data[
                        "signatures"
                    ].get(
                        "인계자 성명",
                        auto_giver
                    ),
                    key="auto_sign_giver"
                )

                auto_sign_giver_date = st.text_input(
                    "인계자 서명일",
                    value=parsed_data[
                        "signatures"
                    ].get(
                        "인계자 서명일",
                        ""
                    ),
                    key="auto_sign_giver_date"
                )

            with s2:
                auto_sign_receiver = st.text_input(
                    "인수자",
                    value=parsed_data[
                        "signatures"
                    ].get(
                        "인수자 성명",
                        auto_receiver
                    ),
                    key="auto_sign_receiver"
                )

                auto_sign_receiver_date = st.text_input(
                    "인수자 서명일",
                    value=parsed_data[
                        "signatures"
                    ].get(
                        "인수자 서명일",
                        ""
                    ),
                    key="auto_sign_receiver_date"
                )

            with s3:
                auto_sign_checker = st.text_input(
                    "확인자(팀장 등)",
                    value=parsed_data[
                        "signatures"
                    ].get(
                        "확인자 성명",
                        ""
                    ),
                    key="auto_sign_checker"
                )

                auto_sign_checker_date = st.text_input(
                    "확인자 서명일",
                    value=parsed_data[
                        "signatures"
                    ].get(
                        "확인자 서명일",
                        ""
                    ),
                    key="auto_sign_checker_date"
                )

            # -----------------------------------------
            # 검토 완료 데이터
            # -----------------------------------------
            reviewed_data = {
                "meta": {
                    "기관명": auto_institution,
                    "부서명": auto_department_meta,
                    "문서번호": auto_document_no,
                    "보존기간": auto_retention,
                },

                "basic": {
                    "소속 부서": auto_department,
                    "직위 / 직책": auto_position,
                    "인계자 성명": auto_giver,
                    "인수자 성명": auto_receiver,
                    "작성일": auto_written_date,
                    "인수인계 완료 예정일": auto_expected_date,
                },

                "tasks": clean_records(
                    auto_tasks_df
                ),

                "details": edited_details,

                "urgent_schedule": clean_records(
                    auto_urgent_df
                ),

                "monthly_schedule": clean_records(
                    auto_monthly_df
                ),

                "communication_note": auto_communication,

                "assets": clean_records(
                    auto_assets_df
                ),

                "checklist": clean_records(
                    auto_checklist_df
                ),

                "signatures": {
                    "인계자 성명": auto_sign_giver,
                    "인계자 서명일": auto_sign_giver_date,
                    "인수자 성명": auto_sign_receiver,
                    "인수자 서명일": auto_sign_receiver_date,
                    "확인자 성명": auto_sign_checker,
                    "확인자 서명일": auto_sign_checker_date,
                },
            }

            st.divider()

            # 준비도
            show_readiness(
                reviewed_data
            )

            st.divider()

            # 최종 다운로드
            st.subheader(
                "📄 최종 인수인계서 생성"
            )

            st.info(
                "위 내용을 확인하거나 수정한 뒤 "
                "아래 버튼으로 최종 문서를 생성하세요."
            )

            final_docx = build_docx(
                reviewed_data
            )

            st.download_button(
                "✅ 검토 완료 - 최종 인수인계서 다운로드",
                data=final_docx,
                file_name="업무_인수인계서_최종.docx",
                mime=(
                    "application/"
                    "vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                use_container_width=True,
                key="auto_final_download"
            )

            with st.expander(
                "🔍 최종 데이터 확인"
            ):
                st.json(
                    reviewed_data
                )

        except Exception as e:
            record_error(
                screen="🤖 업무메모 자동 인수인계",
                function="업무메모 처리 및 문서 생성",
                error=e,
                code="MEMO_PROCESS_ERROR",
                input_value=getattr(uploaded_memo, "name", "업무메모"),
            )
            st.error(
                "업무메모 처리 중 오류가 발생했습니다. 시스템이 자동 감지했습니다."
            )

    else:
        st.caption(
            "업무메모_예시.txt 파일을 업로드해 테스트할 수 있습니다."
        )


# =========================================================
# TAB 3. 직접 작성
# =========================================================
with tab_manual:
    render_floating_error_button(
        screen="✍️ 직접 작성",
        function="인수인계서 직접 작성",
        key_suffix="manual",
    )
    st.header(
        "✍️ 인수인계서 직접 작성"
    )

    st.caption(
        "업무메모 없이도 직접 내용을 입력하여 DOCX를 생성할 수 있습니다."
    )

    with st.expander(
        "문서 기본 정보",
        expanded=True
    ):
        c1, c2, c3, c4 = st.columns(4)

        institution = c1.text_input(
            "기관명",
            key="manual_institution"
        )

        department_meta = c2.text_input(
            "부서명",
            key="manual_department_meta"
        )

        document_no = c3.text_input(
            "문서번호",
            key="manual_document_no"
        )

        retention = c4.text_input(
            "보존기간",
            placeholder="예: 3년 / 폐기 시",
            key="manual_retention"
        )

    st.subheader(
        "I. 기본 정보"
    )

    c1, c2 = st.columns(2)

    with c1:
        department = st.text_input(
            "소속 부서",
            key="manual_department"
        )

        giver = st.text_input(
            "인계자 성명",
            key="manual_giver"
        )

        written_date = st.date_input(
            "작성일",
            value=date.today(),
            key="manual_written_date"
        )

    with c2:
        position = st.text_input(
            "직위 / 직책",
            key="manual_position"
        )

        receiver = st.text_input(
            "인수자 성명",
            key="manual_receiver"
        )

        expected_date = st.date_input(
            "인수인계 완료 예정일",
            value=date.today(),
            key="manual_expected_date"
        )

    st.subheader(
        "II. 담당업무 개요"
    )

    tasks_df = st.data_editor(
        default_rows(
            [
                "주요 업무",
                "업무 목적",
                "업무 프로세스",
                "우선순위",
                "비고",
            ],
            3
        ),
        num_rows="dynamic",
        use_container_width=True,
        key="manual_tasks"
    )

    st.subheader(
        "III. 업무 상세"
    )

    st.caption(
        "MVP에서는 최대 3개 주요 업무를 입력합니다."
    )

    details = []

    tabs = st.tabs(
        [
            "업무 1",
            "업무 2",
            "업무 3",
        ]
    )

    for idx, task_tab in enumerate(
        tabs,
        start=1
    ):
        with task_tab:
            task_name = st.text_input(
                f"업무명 {idx}",
                key=f"manual_task_name_{idx}"
            )

            col1, col2 = st.columns(2)

            with col1:
                overview = st.text_area(
                    "업무 개요",
                    key=f"manual_overview_{idx}"
                )

                kpi = st.text_area(
                    "목적 / 성과지표",
                    key=f"manual_kpi_{idx}"
                )

                project = st.text_area(
                    "진행 중 프로젝트 현황",
                    key=f"manual_project_{idx}"
                )

                regular = st.text_area(
                    "정기 업무",
                    key=f"manual_regular_{idx}"
                )

                irregular = st.text_area(
                    "비정기 업무",
                    key=f"manual_irregular_{idx}"
                )

                deadline = st.text_area(
                    "주요 일정 / 마감",
                    key=f"manual_deadline_{idx}"
                )

                systems = st.text_area(
                    "관련 시스템 / 계정 / 권한",
                    key=f"manual_systems_{idx}"
                )

            with col2:
                contacts = st.text_area(
                    "관련 담당자 / 연락처",
                    key=f"manual_contacts_{idx}"
                )

                collaborate = st.text_area(
                    "협업 부서",
                    key=f"manual_collaborate_{idx}"
                )

                cautions = st.text_area(
                    "특이사항 / 주의사항",
                    key=f"manual_cautions_{idx}"
                )

                risks = st.text_area(
                    "리스크 / 미해결 이슈",
                    key=f"manual_risks_{idx}"
                )

                files = st.text_area(
                    "참고 파일 경로 / 문서 링크",
                    key=f"manual_files_{idx}"
                )

                must_know = st.text_area(
                    "후임자 숙지 필요사항",
                    key=f"manual_must_know_{idx}"
                )

                done = st.selectbox(
                    "인수인계 완료 여부",
                    [
                        "미완료",
                        "진행중",
                        "완료",
                    ],
                    key=f"manual_done_{idx}"
                )

            if any(
                [
                    task_name.strip(),
                    overview.strip(),
                    kpi.strip(),
                    project.strip(),
                    regular.strip(),
                    irregular.strip(),
                    deadline.strip(),
                    systems.strip(),
                    contacts.strip(),
                    collaborate.strip(),
                    cautions.strip(),
                    risks.strip(),
                    files.strip(),
                    must_know.strip(),
                ]
            ):
                details.append(
                    {
                        "업무명": task_name,
                        "업무 개요": overview,
                        "목적 / 성과지표": kpi,
                        "진행 중 프로젝트 현황": project,
                        "정기 업무": regular,
                        "비정기 업무": irregular,
                        "주요 일정 / 마감": deadline,
                        "관련 시스템 / 계정 / 권한": systems,
                        "관련 담당자 / 연락처": contacts,
                        "협업 부서": collaborate,
                        "특이사항 / 주의사항": cautions,
                        "리스크 / 미해결 이슈": risks,
                        "참고 파일 경로 / 문서 링크": files,
                        "후임자 숙지 필요사항": must_know,
                        "인수인계 완료 여부": done,
                    }
                )

    st.subheader(
        "IV. 주요 일정 및 대외 커뮤니케이션"
    )

    st.markdown(
        "#### 1. 긴급 업무 일정"
    )

    urgent_df = st.data_editor(
        default_rows(
            [
                "일자",
                "내용",
                "대응 방법",
                "담당",
            ],
            2
        ),
        num_rows="dynamic",
        use_container_width=True,
        key="manual_urgent"
    )

    st.markdown(
        "#### 2. 향후 1개월 주요 일정"
    )

    monthly_df = st.data_editor(
        default_rows(
            [
                "일자",
                "일정 내용",
                "비고",
            ],
            3
        ),
        num_rows="dynamic",
        use_container_width=True,
        key="manual_monthly"
    )

    st.markdown(
        "#### 3. 대외 커뮤니케이션 유의사항"
    )

    communication_note = st.text_area(
        "커뮤니케이션 유의사항",
        key="manual_communication"
    )

    st.subheader(
        "V. 계정 · 권한 · 자산 인계"
    )

    assets_df = st.data_editor(
        default_rows(
            [
                "시스템 / 자산명",
                "유형",
                "권한 수준",
                "인계 방법",
                "상태",
                "비고",
            ],
            3
        ),
        num_rows="dynamic",
        use_container_width=True,
        key="manual_assets"
    )

    st.subheader(
        "VI. 인수인계 체크리스트"
    )

    check_default = pd.DataFrame(
        [
            {
                "체크": "미완료",
                "항목": "후임자 우선 숙지사항 전달",
                "확인일": "",
                "확인자": "",
                "비고": "",
            },
            {
                "체크": "미완료",
                "항목": "참고 파일 경로 및 링크 전달",
                "확인일": "",
                "확인자": "",
                "비고": "",
            },
            {
                "체크": "미완료",
                "항목": "계정 및 권한 이전 필요사항 전달",
                "확인일": "",
                "확인자": "",
                "비고": "",
            },
            {
                "체크": "미완료",
                "항목": "관련 담당자 및 연락처 전달",
                "확인일": "",
                "확인자": "",
                "비고": "",
            },
        ]
    )

    checklist_df = st.data_editor(
        check_default,
        num_rows="dynamic",
        use_container_width=True,
        key="manual_checklist"
    )

    st.subheader(
        "VII. 서명"
    )

    s1, s2, s3 = st.columns(3)

    with s1:
        sign_giver = st.text_input(
            "인계자 성명",
            value=giver,
            key="manual_sign_giver"
        )

        sign_giver_date = st.date_input(
            "인계자 서명일",
            value=date.today(),
            key="manual_sign_giver_date"
        )

    with s2:
        sign_receiver = st.text_input(
            "인수자 성명",
            value=receiver,
            key="manual_sign_receiver"
        )

        sign_receiver_date = st.date_input(
            "인수자 서명일",
            value=date.today(),
            key="manual_sign_receiver_date"
        )

    with s3:
        sign_checker = st.text_input(
            "확인자(팀장 등) 성명",
            key="manual_sign_checker"
        )

        sign_checker_date = st.date_input(
            "확인자 서명일",
            value=date.today(),
            key="manual_sign_checker_date"
        )

    manual_data = {
        "meta": {
            "기관명": institution,
            "부서명": department_meta,
            "문서번호": document_no,
            "보존기간": retention,
        },

        "basic": {
            "소속 부서": department,
            "직위 / 직책": position,
            "인계자 성명": giver,
            "인수자 성명": receiver,
            "작성일": written_date.isoformat(),
            "인수인계 완료 예정일": expected_date.isoformat(),
        },

        "tasks": clean_records(
            tasks_df
        ),

        "details": details,

        "urgent_schedule": clean_records(
            urgent_df
        ),

        "monthly_schedule": clean_records(
            monthly_df
        ),

        "communication_note": communication_note,

        "assets": clean_records(
            assets_df
        ),

        "checklist": clean_records(
            checklist_df
        ),

        "signatures": {
            "인계자 성명": sign_giver,
            "인계자 서명일": sign_giver_date.isoformat(),
            "인수자 성명": sign_receiver,
            "인수자 서명일": sign_receiver_date.isoformat(),
            "확인자 성명": sign_checker,
            "확인자 서명일": sign_checker_date.isoformat(),
        },
    }

    st.divider()

    show_readiness(
        manual_data
    )

    st.divider()

    st.subheader(
        "📄 자동 생성"
    )

    json_bytes = json.dumps(
        manual_data,
        ensure_ascii=False,
        indent=2
    ).encode(
        "utf-8"
    )

    docx_bytes = build_docx(
        manual_data
    )

    c1, c2 = st.columns(2)

    with c1:
        st.download_button(
            "💾 입력 데이터(JSON) 저장",
            data=json_bytes,
            file_name="handover_data.json",
            mime="application/json",
            use_container_width=True,
            key="manual_json_download"
        )

    with c2:
        st.download_button(
            "📄 인수인계서(DOCX) 생성",
            data=docx_bytes,
            file_name="업무_인수인계서_자동생성.docx",
            mime=(
                "application/"
                "vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
            key="manual_docx_download"
        )

    with st.expander(
        "현재 저장될 데이터 미리보기"
    ):
        st.json(
            manual_data
        )

# =========================================================
# 관리자 오류함
# =========================================================
with tab_error_admin:
    render_floating_error_button(
        screen="🛠 관리자 오류함",
        function="오류 신고 관리",
        key_suffix="error_admin",
    )
    st.header("🛠 관리자 오류함")
    st.caption(
        "시스템이 자동 감지한 오류·응답 지연 정보를 사용자가 확인 후 전송하면 이곳에 접수됩니다. "
        "현재 MVP에서는 JSONL 파일과 현재 세션에 저장됩니다."
    )

    with st.expander("🧪 자동 오류 감지 테스트"):
        st.caption(
            "시연용 버튼입니다. 누르면 시스템이 오류를 자동 감지한 것처럼 리포트 팝업이 자동으로 열립니다."
        )
        if st.button("⚡ 테스트 오류 발생시키기", key="make_demo_error"):
            try:
                raise TimeoutError("시연용: Q&A 응답 시간이 초과되었습니다.")
            except Exception as e:
                record_error(
                    screen="💬 후임자 Q&A",
                    function="후임자 질문 검색",
                    error=e,
                    code="QA_TIMEOUT_DEMO",
                    input_value="A동 보고서는 지금 어디까지 진행됐어?",
                    response_time=8.24,
                )
                st.info(
                    "테스트 오류를 자동 감지했습니다. 오류 리포트 창이 자동으로 열립니다."
                )

    error_reports = list(reversed(st.session_state.get("error_reports", [])))

    if not error_reports:
        st.info("아직 접수된 오류 신고가 없습니다.")
    else:
        report_rows = []
        for report in error_reports:
            response_time = report.get("response_time")
            report_rows.append(
                {
                    "신고번호": report.get("report_id", ""),
                    "신고시각": report.get("reported_at", ""),
                    "화면": report.get("screen", ""),
                    "기능": report.get("function", ""),
                    "에러 코드": report.get("error_code", ""),
                    "응답시간(초)": (
                        round(float(response_time), 2)
                        if response_time is not None
                        else ""
                    ),
                    "상태": report.get("status", "미처리"),
                }
            )

        st.dataframe(
            pd.DataFrame(report_rows),
            use_container_width=True,
            hide_index=True,
        )

        selected_report_id = st.selectbox(
            "상세 확인할 신고",
            [report.get("report_id", "") for report in error_reports],
            key="admin_error_report_select",
        )

        selected_report = next(
            (
                report
                for report in error_reports
                if report.get("report_id") == selected_report_id
            ),
            None,
        )

        if selected_report:
            st.markdown("#### 📋 신고 상세")
            st.write(f"**자동 요약:** {selected_report.get('auto_summary', '')}")

            if selected_report.get("additional_note"):
                st.write(
                    f"**사용자 추가 의견:** {selected_report.get('additional_note', '')}"
                )

            st.text_area(
                "최근 동작 순서",
                value=selected_report.get("action_history", ""),
                height=170,
                disabled=True,
                key=f"admin_action_history_{selected_report_id}",
            )

            with st.expander("기술 정보 보기"):
                st.json(
                    {
                        "error_type": selected_report.get("error_type", ""),
                        "error_message": selected_report.get("error_message", ""),
                        "last_input": selected_report.get("last_input", ""),
                        "response_time": selected_report.get("response_time"),
                        "screenshot": selected_report.get("screenshot", ""),
                        "traceback": selected_report.get("technical_traceback", ""),
                    }
                )


# =========================================================
# 자동 오류 리포트 팝업 트리거
# =========================================================
# record_error()/record_system_issue()가 오류를 감지하면 이 플래그가 켜집니다.
# 사용자는 별도의 신고 아이콘을 누를 필요가 없습니다.
if st.session_state.get("pending_auto_error_dialog"):
    st.session_state["pending_auto_error_dialog"] = False
    st.session_state.pop("error_report_auto_summary", None)
    st.session_state.pop("error_report_additional_note", None)
    show_error_report_dialog()
